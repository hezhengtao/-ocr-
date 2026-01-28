import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import time
import os
import base64
import requests
import json
from datetime import datetime
import re
from PIL import Image, ImageTk, ImageDraw, ImageFont
import sys
import subprocess
import ctypes
import warnings
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import configparser
import difflib  # 新增导入，用于文本相似度匹配

# PyInstaller 打包提示：以下导入仅用于打包时分析依赖
try:
    import cv2
    import numpy as np
except ImportError:
    pass  # 开发环境可能没有安装，打包时 PyInstaller 会分析这些依赖

# 忽略libpng警告
warnings.filterwarnings("ignore", category=UserWarning)

# Windows高DPI自适应
def set_dpi_awareness():
    """设置Windows高DPI感知"""
    try:
        awareness = ctypes.c_int()
        ctypes.windll.shcore.GetProcessDpiAwareness(0, ctypes.byref(awareness))
        ctypes.windll.shcore.SetProcessDpiAwareness(2)
    except:
        try:
            ctypes.windll.user32.SetProcessDPIAware()
        except:
            pass

set_dpi_awareness()

# 多颜色HSV范围定义 - [新增] 任意颜色
COLOR_RANGES = {
    "红色": [
        ((0, 80, 80), (10, 255, 255)),
        ((170, 80, 80), (180, 255, 255))
    ],
    "蓝色": [((100, 80, 80), (130, 255, 255))],
    "绿色": [((40, 80, 80), (80, 255, 255))],
    "紫色": [((125, 80, 80), (150, 255, 255))],
    "黄色": [((20, 80, 80), (35, 255, 255))],
    # [新增] 任意颜色: 排除黑/白/灰 (Saturation > 30, Value > 40)
    "任意颜色": [((0, 30, 40), (180, 255, 255))]
}

class AnnotationDetector:
    """多颜色笔迹检测器"""
    
    @staticmethod
    def detect_annotations(image_path, selected_colors=None, threshold=0.7, merge_distance=20):
        """
        检测图像中的彩色笔迹
        Args:
            image_path: 图像文件路径
            selected_colors: 选择的颜色列表，如["红色", "蓝色"]
            threshold: 检测阈值
            merge_distance: 合并距离（像素）
        Returns:
            检测结果字典
        """
        try:
            if selected_colors is None:
                selected_colors = ["红色"]
            
            # 动态导入cv2和numpy
            import cv2
            import numpy as np
            
            img_array = np.fromfile(image_path, dtype=np.uint8)
            img = cv2.imdecode(img_array, cv2.IMREAD_COLOR)
            
            if img is None:
                raise ValueError(f"无法读取图像: {image_path}")
            
            hsv = cv2.cvtColor(img, cv2.COLOR_BGR2HSV)
            height, width = img.shape[:2]
            
            # 合并所有选中颜色的掩码
            combined_mask = np.zeros((height, width), dtype=np.uint8)
            color_masks = {}
            
            for color_name in selected_colors:
                if color_name in COLOR_RANGES:
                    color_mask = np.zeros((height, width), dtype=np.uint8)
                    for lower, upper in COLOR_RANGES[color_name]:
                        lower_array = np.array(lower, dtype=np.uint8)
                        upper_array = np.array(upper, dtype=np.uint8)
                        mask_part = cv2.inRange(hsv, lower_array, upper_array)
                        color_mask = cv2.bitwise_or(color_mask, mask_part)
                    
                    # 形态学操作
                    kernel = np.ones((3, 3), np.uint8)
                    color_mask = cv2.morphologyEx(color_mask, cv2.MORPH_OPEN, kernel)
                    color_mask = cv2.morphologyEx(color_mask, cv2.MORPH_CLOSE, kernel)
                    
                    if threshold > 0:
                        color_mask = cv2.threshold(color_mask, int(threshold * 255), 255, cv2.THRESH_BINARY)[1]
                    
                    color_masks[color_name] = color_mask
                    combined_mask = cv2.bitwise_or(combined_mask, color_mask)
            
            # 查找所有轮廓
            contours, _ = cv2.findContours(combined_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 合并相近轮廓（关键优化）
            merged_contours = []
            used = [False] * len(contours)
            
            for i in range(len(contours)):
                if used[i]:
                    continue
                
                current_cnt = contours[i]
                current_bbox = cv2.boundingRect(current_cnt)
                
                # 寻找附近轮廓
                merged_points = current_cnt
                for j in range(i+1, len(contours)):
                    if used[j]:
                        continue
                    
                    other_cnt = contours[j]
                    other_bbox = cv2.boundingRect(other_cnt)
                    
                    # 计算轮廓距离
                    dist_x = abs(current_bbox[0] - other_bbox[0])
                    dist_y = abs(current_bbox[1] - other_bbox[1])
                    min_dist = min(dist_x, dist_y)
                    
                    # 如果轮廓在同一行高度范围内且距离较近，合并
                    if (abs(current_bbox[1] - other_bbox[1]) < merge_distance and 
                        min_dist < merge_distance * 3):
                        merged_points = np.concatenate((merged_points, other_cnt))
                        used[j] = True
                
                merged_contours.append(merged_points)
                used[i] = True
            
            # 过滤小区域并识别颜色
            annotations = []
            for cnt in merged_contours:
                area = cv2.contourArea(cnt)
                if area < 50:  # 最小面积阈值
                    continue
                
                x, y, w, h = cv2.boundingRect(cnt)
                
                # 确定笔迹颜色
                color_name = "红色"  # 默认
                max_overlap = 0
                for cname, cmask in color_masks.items():
                    # 计算此轮廓在该颜色掩码中的覆盖比例
                    roi_mask = np.zeros_like(cmask)
                    cv2.drawContours(roi_mask, [cnt], -1, 255, -1)
                    overlap = cv2.countNonZero(cv2.bitwise_and(roi_mask, cmask))
                    if overlap > max_overlap:
                        max_overlap = overlap
                        color_name = cname
                
                annotations.append({
                    'bbox': (x, y, w, h),
                    'color': color_name,
                    'contour': cnt,
                    'area': area,
                    'center': (x + w // 2, y + h // 2)
                })
            
            return {
                'original_image': img,
                'combined_mask': combined_mask,
                'annotations': annotations,
                'color_masks': color_masks,
                'height': height,
                'width': width
            }
            
        except Exception as e:
            return None
    
    @staticmethod
    def extract_dominant_annotation_colors(image_path, max_colors=3):
        """
        提取图像中的主要笔迹颜色
        Args:
            image_path: 图像文件路径
            max_colors: 最大颜色数量
        Returns:
            颜色名称列表
        """
        try:
            # 检测所有颜色的笔迹
            result = AnnotationDetector.detect_annotations(image_path, selected_colors=["红色", "蓝色", "绿色", "紫色", "黄色"])
            if result is None or not result['annotations']:
                return []
            
            # 统计颜色数量
            color_counts = {}
            for ann in result['annotations']:
                color = ann['color']
                color_counts[color] = color_counts.get(color, 0) + 1
            
            # 按数量排序，取前max_colors个
            sorted_colors = sorted(color_counts.items(), key=lambda x: x[1], reverse=True)
            dominant_colors = [color for color, count in sorted_colors[:max_colors]]
            
            return dominant_colors
            
        except Exception as e:
            return []
    
    @staticmethod
    def highlight_annotations(image_path, output_path=None, selected_colors=None):
        """高亮显示检测到的笔迹 - [适配] 支持任意颜色显示"""
        try:
            # 动态导入cv2
            import cv2
            
            result = AnnotationDetector.detect_annotations(image_path, selected_colors)
            if result is None:
                return image_path
            
            img = result['original_image'].copy()
            annotations = result['annotations']
            
            # 定义颜色映射
            color_map = {
                "红色": (0, 0, 255),
                "蓝色": (255, 0, 0),
                "绿色": (0, 255, 0),
                "紫色": (255, 0, 255),
                "黄色": (0, 255, 255),
                "任意颜色": (0, 165, 255) # 橙色用于通用标记
            }
            
            # 在原始图像上绘制轮廓
            for ann in annotations:
                # 如果颜色不在映射中，默认用橙色
                color = color_map.get(ann['color'], (0, 165, 255))
                cv2.drawContours(img, [ann['contour']], -1, color, 2)
                
                # 添加颜色标签
                x, y, w, h = ann['bbox']
                # 防止文字跑出图片
                text_y = y - 5 if y - 5 > 10 else y + h + 15
                cv2.putText(img, ann['color'], (x, text_y), 
                           cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 1)
            
            if output_path is None:
                output_path = image_path.replace('.', '_annotations.')
            
            success, encoded_img = cv2.imencode('.jpg', img)
            if success:
                encoded_img.tofile(output_path)
            
            return output_path
            
        except Exception as e:
            return image_path
    
    @staticmethod
    def ocr_annotation_region(image, mask, bbox, token, api_url="https://n8q0m2jaw0j292wf.aistudio-app.com/ocr"):
        """[手写专用版] 批注OCR：手写模型优先 + 笔画加粗 + 强力清洗"""
        try:
            import cv2
            import numpy as np
            import requests
            import re
            
            x, y, w, h = bbox
            
            # 1. 扩大截图范围 (防止字被切断)
            padding = 15
            h_img, w_img = image.shape[:2]
            x1 = max(0, x - padding)
            y1 = max(0, y - padding)
            x2 = min(w_img, x + w + padding)
            y2 = min(h_img, y + h + padding)
            
            crop_img = image[y1:y2, x1:x2]
            crop_mask = mask[y1:y2, x1:x2]
            
            if crop_img.size == 0: return ""

            # ================= 图像增强：专为手写优化 =================
            
            # 1. 颜色提纯 (去除非红色的印刷字干扰)
            white_bg = np.ones_like(crop_img) * 255
            masked_img = cv2.bitwise_and(crop_img, crop_img, mask=crop_mask)
            bg_mask = cv2.bitwise_not(crop_mask)
            white_bg_part = cv2.bitwise_and(white_bg, white_bg, mask=bg_mask)
            final_img = cv2.add(masked_img, white_bg_part)
            
            # 2. 转灰度并锐化 (突出笔锋)
            gray = cv2.cvtColor(final_img, cv2.COLOR_BGR2GRAY)
            kernel_sharpen = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
            sharpened = cv2.filter2D(gray, -1, kernel_sharpen)
            
            # 3. 笔画加粗 (关键：解决红笔字太细的问题)
            # 二值化
            _, binary = cv2.threshold(sharpened, 200, 255, cv2.THRESH_BINARY)
            # 反转为白字黑底
            inverted = cv2.bitwise_not(binary)
            # 膨胀 (加粗)
            kernel_dilate = np.ones((2, 2), np.uint8)
            dilated = cv2.dilate(inverted, kernel_dilate, iterations=1)
            # 转回白底黑字
            thick_img = cv2.bitwise_not(dilated)
            
            # 4. 智能放大
            h_crop, w_crop = thick_img.shape[:2]
            if h_crop < 64: 
                scale = 64 / h_crop
                thick_img = cv2.resize(thick_img, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

            _, encoded_img = cv2.imencode('.jpg', thick_img)
            b64_data = base64.b64encode(encoded_img).decode('utf-8')

            # ================= OCR 识别逻辑 =================
            def call_ocr(model_type):
                try:
                    payload = {
                        "file": b64_data,
                        "fileType": 1,
                        "use_doc_preprocessor": False,
                        "text_type": model_type, # 动态切换模型
                        "rec_model_type": "ch"
                    }
                    # 超时设置短一点，以便快速重试
                    resp = requests.post(api_url, json=payload, headers={"Authorization": f"token {token}"}, timeout=10)
                    
                    if resp.status_code == 200:
                        data = resp.json().get('result', {})
                        texts = []
                        
                        # V5 嵌套结构解析 (防止出现 modelsettings 乱码)
                        if isinstance(data, dict):
                            if 'ocrResults' in data:
                                for item in data['ocrResults']:
                                    p = item.get('prunedResult', item)
                                    if isinstance(p, dict) and 'rec_texts' in p:
                                        texts.extend(p['rec_texts'])
                                    elif 'text' in item:
                                        texts.append(item['text'])
                            elif 'rec_texts' in data:
                                texts.extend(data['rec_texts'])
                        
                        return "".join([str(t) for t in texts if t])
                    return ""
                except:
                    return ""

            # 策略：优先用"手写模型"，如果结果为空或太短，用"通用模型"兜底
            final_text = call_ocr("handwriting")
            
            # 如果手写模型没认出来，或者认出来的是乱码，尝试通用模型
            if not final_text or len(final_text) < 2:
                general_text = call_ocr("general")
                if len(general_text) > len(final_text):
                    final_text = general_text

            # ================= 后处理与纠错 =================
            
            # 1. 强力过滤 API 配置乱码 (彻底解决 modelsettings 问题)
            if "modelsettings" in final_text.lower() or "docpreprocessor" in final_text.lower():
                return ""

            # 2. 常见错别字纠错 (针对手写体)
            CORRECTION_MAP = {
                "改柄句": "改病句", "改病奇": "改病句", "改痛句": "改病句", "柄句": "病句",
                "答配不当": "搭配不当", "搭西已不当": "搭配不当", "塔配不当": "搭配不当", 
                "配配不": "搭配不当", "已不当": "配不当",
                "前后矛有": "前后矛盾", "前后矛后": "前后矛盾", "前后矛": "前后矛盾",
                "早经": "本草经", "元": "无", "已": "已知",
                "sumra": "", "Fnonmmly": "" # 过滤常见英文乱码
            }
            
            for wrong, right in CORRECTION_MAP.items():
                if wrong in final_text:
                    final_text = final_text.replace(wrong, right)
            
            # 3. 最终清洗 (保留中文、英文、数字、关键符号)
            # 允许 A-D (选项), 0-9, 中文, 常见标点
            clean_text = re.sub(r'[^\u4e00-\u9fffA-Da-d0-9\(\)（）,.?!，。？！%√×]', '', final_text)
            
            # 4. 过滤单字符误读 (除了 A-D 和 对勾)
            if len(clean_text) == 1:
                if clean_text in ["0", "o", "O", ",", ".", "-"]: return ""
                if clean_text not in ["A", "B", "C", "D", "√", "×"] and not re.match(r'[\u4e00-\u9fff]', clean_text):
                    return ""

            return clean_text
            
        except Exception as e:
            print(f"Annotation OCR Error: {e}")
            return ""

class Tooltip:
    """✅ 修复：全局Tooltip类 - 防止闪烁，创建一次显示/隐藏"""
    def __init__(self, widget, text):
        self.widget = widget
        self.text = text
        self.tip = None
        self.show_timer = None
        self.hide_timer = None
        self.is_visible = False
        
        widget.bind("<Enter>", self.schedule_show)
        widget.bind("<Leave>", self.schedule_hide)
        widget.bind("<ButtonPress>", self.hide)
    
    def schedule_show(self, event=None):
        """延迟显示Tooltip"""
        self.unschedule()
        self.show_timer = self.widget.after(500, self.show)
    
    def schedule_hide(self, event=None):
        """延迟隐藏Tooltip"""
        self.unschedule()
        self.hide_timer = self.widget.after(300, self.hide)
    
    def unschedule(self):
        """取消定时器"""
        if self.show_timer:
            self.widget.after_cancel(self.show_timer)
            self.show_timer = None
        if self.hide_timer:
            self.widget.after_cancel(self.hide_timer)
            self.hide_timer = None
    
    def show(self, event=None):
        """✅ 修复：Tooltip显示位置计算，考虑高DPI缩放"""
        if self.is_visible:
            return
            
        # [修复] 查找 App 实例
        try:
            app = self.widget.winfo_toplevel().app
            theme_mode = app.theme_mode
            scale_factor = app.scale_factor
        except AttributeError:
            theme_mode = "light"
            scale_factor = 1.0
        
        # 颜色设置
        bg_color = "#111827" if theme_mode == "dark" else "#f8fafc"
        fg_color = "#ffffff" if theme_mode == "dark" else "#0f172a"
        
        # ✅ 关键修复：使用 winfo_pointerxy 获取当前鼠标位置
        x, y = self.widget.winfo_pointerxy()
        
        # 偏移量，让 tooltip 显示在鼠标右下方
        offset_x = int(10 * scale_factor)
        offset_y = int(10 * scale_factor)
        
        # 防止超出屏幕右侧
        screen_width = self.widget.winfo_screenwidth()
        tip_width = len(self.text) * int(6 * scale_factor) + int(16 * scale_factor)
        if x + tip_width + offset_x > screen_width:
            x = screen_width - tip_width - offset_x
        
        # 防止超出屏幕底部
        screen_height = self.widget.winfo_screenheight()
        tip_height = int(30 * scale_factor)
        if y + tip_height + offset_y > screen_height:
            y = self.widget.winfo_rooty() - tip_height - offset_y
        else:
            y = y + offset_y

        if self.tip is None:
            self.tip = tk.Toplevel(self.widget)
            self.tip.wm_overrideredirect(True)
            
            self.tip_label = tk.Label(
                self.tip,
                text=self.text,
                bg=bg_color,
                fg=fg_color,
                font=("Microsoft YaHei UI", int(9 * scale_factor)),
                padx=int(8 * scale_factor),
                pady=int(6 * scale_factor),
                relief="solid",
                borderwidth=1
            )
            self.tip_label.pack()
        
        # 更新位置和样式
        self.tip.wm_geometry(f"+{x}+{y}")
        self.tip_label.config(
            bg=bg_color,
            fg=fg_color,
            font=("Microsoft YaHei UI", int(9 * scale_factor))
        )
        
        # 显示
        self.tip.deiconify()
        self.tip.lift()
        self.is_visible = True
        
        # 淡入效果
        self.tip.attributes('-alpha', 0.0)
        for i in range(1, 11):
            self.tip.attributes('-alpha', i * 0.1)
            self.tip.update()
            time.sleep(0.01)
    
    def hide(self, event=None):
        """隐藏Tooltip"""
        if self.tip and self.is_visible:
            # 淡出效果
            for i in range(9, -1, -1):
                if self.tip:
                    self.tip.attributes('-alpha', i * 0.1)
                    self.tip.update()
                    time.sleep(0.01)
            self.tip.withdraw()
            self.is_visible = False
        self.unschedule()

class ModernCheckbutton:
    """现代化复选框 - [实时渲染版] 解决一切背景色同步延迟"""
    def __init__(self, parent, text="", variable=None, command=None, **kwargs):
        self.parent = parent
        self.text = text
        self.variable = variable
        self.command = command
        
        self.frame = tk.Frame(parent)
        
        self.scale_factor = 1.0
        try:
            self.scale_factor = parent.winfo_fpixels('1i') / 96.0
        except: pass

        self.size = int(20 * self.scale_factor)
        
        self.canvas = tk.Canvas(self.frame, width=self.size, height=self.size, 
                                highlightthickness=0, bd=0)
        self.canvas.pack(side=tk.LEFT)
        
        self.label = tk.Label(self.frame, text=text, font=("Microsoft YaHei UI", int(9 * self.scale_factor)))
        self.label.pack(side=tk.LEFT, padx=(int(6 * self.scale_factor), 0))
        
        self.state = tk.NORMAL
        self.selected = False
        
        if variable:
            self.selected = variable.get()
            self.variable.trace_add("write", lambda *args: self._update_from_var())
        
        self.canvas.bind("<Button-1>", self._toggle)
        self.label.bind("<Button-1>", self._toggle)
        
        self._register_for_refresh()
        
        # 立即更新一次
        self.update_theme()

    def _register_for_refresh(self):
        app = self._get_app()
        if app and hasattr(app, 'register_refresh_widget'):
            app.register_refresh_widget(self)

    def _get_app(self):
        try:
            return self.parent.winfo_toplevel().app
        except AttributeError:
            return None

    def update_theme(self):
        """更新主题 - 仅设置属性，绘制逻辑后移"""
        app = self._get_app()
        if not app: return
        
        # 触发重绘，重绘时会去拿最新的颜色
        self._draw()

    def _draw(self):
        """实时获取颜色并绘制"""
        app = self._get_app()
        if not app: return

        mode = app.theme_mode
        colors = app.colors[mode]
        is_dark = (mode == "dark")
        
        # === [绝杀] 强制背景色逻辑 ===
        # 你的复选框都在 LabelFrame (Card) 里，所以背景色必须是 card 色。
        target_bg = colors["card"]

        # 实时设置控件颜色
        self.frame.configure(bg=target_bg)
        self.canvas.configure(bg=target_bg)
        self.label.configure(bg=target_bg, fg=colors["text"])
        
        # 准备绘制参数
        if is_dark:
            box_bg = target_bg           # 透底
            box_selected_bg = "#00FF7F"  # 荧光绿
            border_color = "#00FF7F"
            tick_color = "#000000"       # 黑色对勾
            line_width = 1
        else:
            box_bg = "#ffffff"           # 白底
            box_selected_bg = colors["primary"]
            border_color = "#cbd5e1"
            tick_color = "#FFFFFF"
            line_width = 2

        self.canvas.delete("all")
        pad = 2 
        
        if self.selected:
            # 绘制选中状态
            self.canvas.create_rectangle(
                pad, pad, self.size-pad, self.size-pad,
                fill=box_selected_bg, 
                outline=border_color, 
                width=1
            )
            # 对勾
            points = [
                (self.size * 0.25, self.size * 0.5),
                (self.size * 0.45, self.size * 0.75),
                (self.size * 0.8,  self.size * 0.3)
            ]
            self.canvas.create_line(
                points, 
                fill=tick_color, 
                width=int(2 * self.scale_factor),
                capstyle="round", 
                joinstyle="round"
            )
        else:
            # 绘制未选中状态
            self.canvas.create_rectangle(
                pad, pad, self.size-pad, self.size-pad,
                fill=box_bg, 
                outline=border_color, 
                width=line_width
            )

    def _toggle(self, event=None):
        if self.state == tk.DISABLED: return
        self.selected = not self.selected
        if self.variable: self.variable.set(self.selected)
        self._draw()
        if self.command: self.command()

    def _update_from_var(self):
        if self.variable:
            val = self.variable.get()
            if self.selected != val:
                self.selected = val
                self._draw()

    def grid(self, **kwargs): self.frame.grid(**kwargs)
    def pack(self, **kwargs): self.frame.pack(**kwargs)
    def place(self, **kwargs): self.frame.place(**kwargs)

class ModernButton(tk.Button):
    """现代化按钮 - [最终完美修复版]"""
    def __init__(self, parent, **kwargs):
        self.variant = kwargs.pop("variant", "primary")
        self.parent = parent
        
        self.scale_factor = 1.0
        try:
            self.scale_factor = parent.winfo_fpixels('1i') / 96.0
        except: pass
        
        super().__init__(parent, **kwargs)
        
        self.configure(
            relief="flat", bd=0, cursor="hand2",
            font=("Microsoft YaHei UI", int(9 * self.scale_factor))
        )
        
        self.bind("<Enter>", self._on_hover)
        self.bind("<Leave>", self._on_leave)
        
        self._register_for_refresh()
        self.update_theme()

    def _register_for_refresh(self):
        app = self._get_app()
        if app and hasattr(app, 'register_refresh_widget'):
            app.register_refresh_widget(self)
            
    def _get_app(self):
        try:
            return self.parent.winfo_toplevel().app
        except AttributeError:
            return None

    def update_theme(self):
        app = self._get_app()
        if not app: return
        
        mode = app.theme_mode
        
        if mode == "dark":
            if self.variant == "primary":
                self.normal_bg = "#006400"
                self.normal_fg = "#FFFFFF"
                self.hover_bg = "#008000"
            elif self.variant == "danger":
                self.normal_bg = "#8B0000"
                self.normal_fg = "#FFFFFF"
                self.hover_bg = "#FF0000"
            else:
                self.normal_bg = "#333333"
                self.normal_fg = "#00FF7F"
                self.hover_bg = "#444444"
        else:
            if self.variant == "primary":
                self.normal_bg = "#3b82f6"
                self.normal_fg = "#FFFFFF"
                self.hover_bg = "#2563eb"
            elif self.variant == "danger":
                self.normal_bg = "#ef4444"
                self.normal_fg = "#FFFFFF"
                self.hover_bg = "#dc2626"
            else:
                self.normal_bg = "#e2e8f0"
                self.normal_fg = "#0f172a"
                self.hover_bg = "#cbd5e1"
        
        self.configure(bg=self.normal_bg, fg=self.normal_fg, activebackground=self.hover_bg, activeforeground=self.normal_fg)

    def _on_hover(self, e):
        if self['state'] != 'disabled' and hasattr(self, 'hover_bg'):
            self.configure(bg=self.hover_bg)

    def _on_leave(self, e):
        if self['state'] != 'disabled' and hasattr(self, 'normal_bg'):
            self.configure(bg=self.normal_bg)

class OCRDesktopApp:
    def __init__(self, root):
        self.root = root
        self.root.app = self
        
        self.root.title("智能OCR处理器-by HZT")
        
        # === [修复 1] Logo 加载与持久化 ===
        # 必须保存为 self.app_icon，供后续弹窗使用
        self.app_icon = None 
        try:
            # 尝试加载 PNG (推荐)
            png_path = self.get_resource_path("logo.png")
            ico_path = self.get_resource_path("logo.ico")
            
            # 设置 Windows 任务栏 ID
            import ctypes
            myappid = 'mycompany.ocr.desktop.final.v6' 
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)

            if os.path.exists(png_path):
                from PIL import Image, ImageTk
                img = Image.open(png_path).convert("RGBA")
                self.app_icon = ImageTk.PhotoImage(img) 
                # 设置主窗口图标 (True 表示传递给子窗口，但在 Windows 上往往无效，需要手动设)
                self.root.iconphoto(True, self.app_icon)
            elif os.path.exists(ico_path):
                self.root.iconbitmap(ico_path)
        except Exception as e:
            print(f"Icon load error: {e}")
        
        # ... (以下代码保持不变) ...
        self.base_font_size = 8
        self.title_font_size = 9
        self.small_font_size = 7
        
        self.scale_factor = root.winfo_fpixels('1i') / 96.0
        root.tk.call('tk', 'scaling', self.scale_factor)
        
        screen_width = root.winfo_screenwidth()
        screen_height = root.winfo_screenheight()
        window_width = int(1200 * self.scale_factor)
        window_height = int(700 * self.scale_factor)
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        
        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        self.root.minsize(int(900 * self.scale_factor), int(500 * self.scale_factor))
        self.root.configure(bg="#ffffff")
        
        self.token = self.load_api_token()
        self.theme_names = {"light": "浅色模式", "dark": "深色模式"}
        
        self.files = []
        self.file_previews = {}
        self.processing = False
        self.output_prefix = ""
        
        self.font_family = "Microsoft YaHei UI"
        self.base_font = (self.font_family, int(self.base_font_size * self.scale_factor))
        self.title_font = (self.font_family, int(self.title_font_size * self.scale_factor))
        self.small_font = (self.font_family, int(self.small_font_size * self.scale_factor))
        
        self.theme_mode = "light"
        self.disable_tooltips = False
        self.refreshables = []
        
        self.selected_colors = {
            "红色": tk.BooleanVar(value=True),
            "蓝色": tk.BooleanVar(value=False),
            "绿色": tk.BooleanVar(value=False),
            "紫色": tk.BooleanVar(value=False),
            "黄色": tk.BooleanVar(value=False),
            "任意颜色": tk.BooleanVar(value=False),
            "自定义颜色1": tk.BooleanVar(value=False),
            "自定义颜色2": tk.BooleanVar(value=False)
        }
        
        self.custom_colors = {"自定义颜色1": "#FF0000", "自定义颜色2": "#0000FF"}
        
        self.prefix_var = tk.StringVar()
        self.output_path_var = tk.StringVar()
        
        self.enable_question_detection = tk.BooleanVar(value=True)
        self.enable_option_detection = tk.BooleanVar(value=True)
        self.enable_annotation_detection = tk.BooleanVar(value=False)
        self.output_format_mode = tk.StringVar(value="自动分行")
        self.color_tolerance = tk.IntVar(value=30)
        
        # 修改这里：强制开启自动排序
        self.sort_by_question_var = tk.BooleanVar(value=True)
        
        self.question_number_formats = {
            "1.": tk.BooleanVar(value=True),
            "1、": tk.BooleanVar(value=True),
            "(1)": tk.BooleanVar(value=True),
            "一、": tk.BooleanVar(value=True),
            "①②③": tk.BooleanVar(value=True),
        }
        
        self.check_dependencies()
        self.init_configs()
        
        self.api_var = tk.StringVar(value="PaddleOCR-VL")
        self.format_var = tk.StringVar(value="Word (.docx)")
        self.merge_var = tk.BooleanVar(value=True)
        
        self.setup_styles()
        self.setup_ui()
        self.setup_shortcuts()
        
        self.theme_mode = "light"
        self.switch_theme()
        
        self.root.after(100, lambda: self.set_theme("light")) 
        
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def get_app_path(self):
        import sys
        if getattr(sys, 'frozen', False):
            return os.path.dirname(sys.executable)
        else:
            return os.path.dirname(os.path.abspath(__file__))
    
    def get_resource_path(self, relative_path):
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")
        
        return os.path.join(base_path, relative_path)
    
    def load_api_token(self):
        token = os.environ.get("OCR_API_TOKEN")
        if token: 
            return token
        
        config_file = os.path.join(self.get_app_path(), "config.ini")
        if os.path.exists(config_file):
            try:
                config = configparser.ConfigParser()
                config.read(config_file, encoding='utf-8')
                if 'API' in config and 'token' in config['API']:
                    t = config['API']['token']
                    if t and len(t) > 10 and "replace" not in t.lower():
                        return t
            except:
                pass
        
        return ""
    
    def create_config_template(self, config_file, default_token):
        config = configparser.ConfigParser()
        config['API'] = {
            'token': '',
            '# 说明': '请填写您的 API Token',
            '# 获取方式': '访问 https://aistudio.baidu.com 创建应用获取'
        }
        config['Settings'] = {
            'output_path': 'OCR_Output',
            'default_format': 'docx'
        }
        
        with open(config_file, 'w', encoding='utf-8') as f:
            config.write(f)
        
        self.log_message("✅ 已创建配置文件模板: config.ini", "信息")
        self.log_message("⚠️ 请编辑此文件并填入您自己的 API Token", "警告")
    
    def register_refresh_widget(self, widget):
        if widget not in self.refreshables:
            self.refreshables.append(widget)
    
    def check_dependencies(self):
        try:
            import cv2
            import numpy as np
            self.has_opencv = True
            self.log_message("✅ OpenCV 已安装，笔迹检测功能可用", "信息")
        except ImportError:
            self.has_opencv = False
            self.log_message("⚠️ OpenCV 未安装，笔迹检测功能不可用", "警告")

    def init_configs(self):
        self.api_configs = {
            "PaddleOCR-VL": {"url": "https://iaj1g3i2s5m3w9l9.aistudio-app.com/layout-parsing", "description": "复杂图文解析", "color": "#3b82f6"},
            "PP-OCRv5": {"url": "https://n8q0m2jaw0j292wf.aistudio-app.com/ocr", "description": "高速纯文字识别", "color": "#10b981"},
            "PP-StructureV3": {"url": "https://g5l6b4cav1f5g3jb.aistudio-app.com/layout-parsing", "description": "通用文档还原", "color": "#f59e0b"}
        }
        self.output_formats = {"Markdown (.md)": "md", "纯文本 (.txt)": "txt", "JSON (.json)": "json", "Word (.docx)": "docx"}
        
        self.write_annotations_to_doc = tk.BooleanVar(value=True)

    def setup_styles(self):
        style = ttk.Style()
        style.theme_use('clam')
        
        self.colors = {
            "light": {
                "bg": "#ffffff",
                "card": "#ffffff",
                "text": "#000000",
                "subtext": "#475569",
                "border": "#e2e8f0",
                "primary": "#3b82f6",
                "primary_hover": "#2563eb",
                "secondary": "#94a3b8",
                "success": "#10b981",
                "warning": "#facc15",
                "danger": "#ef4444",
                "input_bg": "#f1f5f9",
                "select": "#3b82f6",
                "button_text": "white",
                "tab_bg": "#ffffff",
                "tab_active": "#3b82f6",
                "tab_inactive": "#64748b",
                "log_text": "#475569"
            },
            "dark": {
                "bg": "#121212",
                "card": "#1E1E1E",
                "text": "#00FF7F",
                "subtext": "#90EE90",
                "border": "#333333",
                "primary": "#006400",
                "primary_hover": "#008000",
                "secondary": "#2F4F4F",
                "success": "#00FF00",
                "warning": "#FFD700",
                "danger": "#DC143C",
                "input_bg": "#2D2D2D",
                "select": "#00FF7F",
                "button_text": "#FFFFFF",
                "tab_bg": "#1E1E1E",
                "tab_active": "#00FF7F",
                "tab_inactive": "#556B2F",
                "log_text": "#00FF7F"
            }
        }
        
        style.configure("TFrame", background=self.colors["light"]["bg"])
        style.configure("Card.TFrame", background=self.colors["light"]["card"])
        
        # === [核心修复] 初始化 Notebook 样式 ===
        # 显式设置 lightcolor/darkcolor 为背景色，消除 3D 白边
        bg_color = self.colors["light"]["bg"]
        style.configure("TNotebook", 
                       background=bg_color, 
                       borderwidth=0,
                       lightcolor=bg_color, # 消除亮边
                       darkcolor=bg_color)  # 消除暗边
                       
        style.configure("TNotebook.Tab", 
                       font=self.title_font,
                       background=self.colors["light"]["tab_bg"],
                       foreground=self.colors["light"]["tab_inactive"],
                       padding=[int(8 * self.scale_factor), int(4 * self.scale_factor)],
                       borderwidth=0)
                       
        style.map("TNotebook.Tab",
                 background=[("selected", self.colors["light"]["bg"])],
                 foreground=[("selected", self.colors["light"]["tab_active"])])
        
        self.root.option_add("*TCombobox*Listbox.font", self.base_font)
        
        style.configure("TButton",
                       padding=[int(5 * self.scale_factor), int(3 * self.scale_factor)])

    def setup_ui(self):
        self.main_container = tk.Frame(self.root, bg=self.colors["light"]["bg"])
        self.main_container.pack(fill=tk.BOTH, expand=True, padx=int(10 * self.scale_factor), pady=int(10 * self.scale_factor))
        
        self.main_paned = ttk.PanedWindow(self.main_container, orient=tk.HORIZONTAL)
        self.main_paned.pack(fill=tk.BOTH, expand=True)
        
        self.setup_left_panel()
        self.setup_right_panel()

    def setup_left_panel(self):
        # === [修复] 左侧面板边框 ===
        # 使用 relief="flat" + highlightthickness=1，确保边框颜色完全受控
        # 避免 relief="solid" 产生的系统默认亮色边框
        self.left_panel = tk.Frame(self.main_paned, bg=self.colors["light"]["card"],
                                  highlightthickness=1, 
                                  highlightbackground=self.colors["light"]["border"],
                                  relief="flat") # 改为flat
        self.main_paned.add(self.left_panel, weight=1)
        
        toolbar = tk.Frame(self.left_panel, bg=self.colors["light"]["card"])
        toolbar.pack(fill=tk.X, padx=int(10 * self.scale_factor), pady=int(10 * self.scale_factor))
        
        add_file_btn = ModernButton(toolbar, text="📁 添加文件", variant="primary", 
                                   command=self.select_images, 
                                   padx=int(4 * self.scale_factor),
                                   pady=int(2 * self.scale_factor),
                                   font=self.base_font)
        add_file_btn.pack(side=tk.LEFT, padx=(0, int(5 * self.scale_factor)))
        Tooltip(add_file_btn, "选择要识别的图片或PDF文件")
        
        add_folder_btn = ModernButton(toolbar, text="📂 添加文件夹", variant="secondary",
                                     command=self.select_folder,
                                     padx=int(4 * self.scale_factor),
                                     pady=int(2 * self.scale_factor),
                                     font=self.base_font)
        add_folder_btn.pack(side=tk.LEFT, padx=(0, int(5 * self.scale_factor)))
        Tooltip(add_folder_btn, "选择包含图片/PDF的文件夹")
        
        clear_btn = ModernButton(toolbar, text="清空列表", variant="secondary",
                                command=self.clear_list,
                                padx=int(4 * self.scale_factor),
                                pady=int(2 * self.scale_factor),
                                font=self.base_font)
        clear_btn.pack(side=tk.RIGHT)
        Tooltip(clear_btn, "清空文件列表")
        
        preview_container = tk.Frame(self.left_panel, bg=self.colors["light"]["card"])
        preview_container.pack(fill=tk.BOTH, expand=True, padx=int(10 * self.scale_factor), pady=(0, int(10 * self.scale_factor)))
        
        preview_container.grid_rowconfigure(0, weight=1)
        preview_container.grid_columnconfigure(0, weight=1)
        
        self.preview_canvas = tk.Canvas(preview_container, bg=self.colors["light"]["card"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(preview_container, orient=tk.VERTICAL, command=self.preview_canvas.yview)
        self.preview_canvas.configure(yscrollcommand=scrollbar.set)
        
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.preview_canvas.grid(row=0, column=0, sticky="nsew")
        
        self.preview_content = tk.Frame(self.preview_canvas, bg=self.colors["light"]["card"])
        self.preview_canvas.create_window((0, 0), window=self.preview_content, anchor=tk.NW)
        
        self.preview_content.bind("<Configure>", 
                                 lambda e: self.preview_canvas.configure(
                                     scrollregion=self.preview_canvas.bbox("all")))
        self.preview_canvas.bind_all("<MouseWheel>", self.on_mousewheel)

    def setup_right_panel(self):
        self.right_panel = tk.Frame(self.main_paned, bg=self.colors["light"]["bg"])
        self.main_paned.add(self.right_panel, weight=2)
        
        self.right_panel.grid_rowconfigure(0, weight=1)
        self.right_panel.grid_columnconfigure(0, weight=1)
        
        self.right_content = tk.Frame(self.right_panel, bg=self.colors["light"]["bg"])
        self.right_content.grid(row=0, column=0, sticky="nsew")
        
        self.right_content.grid_rowconfigure(0, weight=1)
        self.right_content.grid_rowconfigure(1, weight=0)
        self.right_content.grid_columnconfigure(0, weight=1)
        
        self.notebook = ttk.Notebook(self.right_content)
        self.notebook.grid(row=0, column=0, sticky="nsew", pady=(0, int(5 * self.scale_factor)))
        
        self.create_tab("⚙️ 配置", self.setup_basic_tab)
        self.create_tab("🔍 识别", self.setup_recognition_tab)
        self.create_tab("📤 输出", self.setup_output_tab)
        self.create_tab("📜 日志", self.setup_log_tab)
        
        self.setup_control_bar()
        
        self.progress_container = tk.Frame(self.right_content, bg=self.colors["light"]["bg"], height=0)
        self.progress_container.grid(row=2, column=0, sticky="ew", pady=(0, 0))
        self.progress_container.grid_propagate(False)
        
        self.progress_bar = ttk.Progressbar(
            self.progress_container, 
            mode="indeterminate",
            length=260
        )

    def setup_control_bar(self):
        self.control_bar = tk.Frame(self.right_content, bg=self.colors["light"]["card"],
                              height=int(45 * self.scale_factor),
                              highlightthickness=1, highlightbackground=self.colors["light"]["border"])
        self.control_bar.grid(row=1, column=0, sticky="ew", pady=(0, 0))
        self.control_bar.grid_propagate(False)
        
        self.control_bar.grid_rowconfigure(0, weight=1)
        self.control_bar.grid_columnconfigure(0, weight=1)
        
        btn_container = tk.Frame(self.control_bar, bg=self.colors["light"]["card"])
        btn_container.place(relx=0, rely=0, relwidth=1, relheight=1)
        
        left_btn_frame = tk.Frame(btn_container, bg=self.colors["light"]["card"])
        left_btn_frame.pack(side=tk.LEFT, padx=int(10 * self.scale_factor))
        
        open_folder_btn = ModernButton(left_btn_frame, text="📂 打开输出目录", variant="secondary",
                                      command=self.open_output_folder, 
                                      padx=int(4 * self.scale_factor),
                                      pady=int(2 * self.scale_factor),
                                      font=self.base_font)
        open_folder_btn.pack(side=tk.LEFT)
        Tooltip(open_folder_btn, "打开保存结果的文件夹")
        
        if self.has_opencv:
            preview_btn = ModernButton(left_btn_frame, text="🔍 预览批注检测", variant="secondary",
                                      command=self.preview_annotation_detection,
                                      padx=int(4 * self.scale_factor),
                                      pady=int(2 * self.scale_factor),
                                      font=self.base_font)
            preview_btn.pack(side=tk.LEFT, padx=(int(5 * self.scale_factor), 0))
            Tooltip(preview_btn, "预览批注检测结果")
        
        right_btn_frame = tk.Frame(btn_container, bg=self.colors["light"]["card"])
        right_btn_frame.pack(side=tk.RIGHT, padx=int(10 * self.scale_factor))
        
        self.stop_button = ModernButton(right_btn_frame, text="⏹ 停止", variant="danger",
                                       command=self.stop_processing, state=tk.DISABLED,
                                       padx=int(4 * self.scale_factor),
                                       pady=int(2 * self.scale_factor),
                                       font=self.base_font)
        self.stop_button.pack(side=tk.RIGHT, padx=(int(5 * self.scale_factor), 0))
        Tooltip(self.stop_button, "停止当前处理任务")
        
        self.start_button = ModernButton(right_btn_frame, text="▶ 开始处理 (F5)", variant="primary",
                                        command=self.start_processing,
                                        padx=int(5 * self.scale_factor),
                                        pady=int(2 * self.scale_factor),
                                        font=self.base_font)
        self.start_button.pack(side=tk.RIGHT)
        Tooltip(self.start_button, "开始处理所有文件 (快捷键: F5)")

    def create_tab(self, title, func):
        # [修复] 强制去边框：bd=0, highlightthickness=0
        # 防止 Frame 组件自带的 1px 默认边框在深色下显形
        tab_frame = tk.Frame(self.notebook, bg=self.colors["light"]["card"], bd=0, highlightthickness=0)
        self.notebook.add(tab_frame, text=title)
        func(tab_frame)

    def setup_basic_tab(self, parent):
        parent.configure(bg=self.colors["light"]["card"])
        
        canvas = tk.Canvas(parent, bg=self.colors["light"]["card"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.colors["light"]["card"])
        
        frame_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def _configure_width(event):
            canvas.itemconfig(frame_id, width=event.width)
        
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", _configure_width)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        PAD_X = int(15 * self.scale_factor)
        PAD_Y = int(20 * self.scale_factor)
        
        theme_frame = tk.LabelFrame(scrollable_frame, text=" 🎨 界面外观 ", 
                                   bg=self.colors["light"]["card"], fg=self.colors["light"]["subtext"],
                                   font=self.base_font, relief="groove", bd=1)
        theme_frame.pack(fill=tk.X, padx=PAD_X, pady=(PAD_Y, 0), ipady=5)
        
        btn_frame = tk.Frame(theme_frame, bg=self.colors["light"]["card"])
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        self.light_btn = ModernButton(btn_frame, text="🌞 浅色模式", 
                                     command=lambda: self.set_theme("light"),
                                     variant="primary" if self.theme_mode == "light" else "secondary",
                                     padx=15, pady=6)
        self.light_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        self.dark_btn = ModernButton(btn_frame, text="🌙 深色模式", 
                                    command=lambda: self.set_theme("dark"),
                                    variant="primary" if self.theme_mode == "dark" else "secondary",
                                    padx=15, pady=6)
        self.dark_btn.pack(side=tk.LEFT)

        api_frame = tk.LabelFrame(scrollable_frame, text=" 🚀 识别引擎与凭证 ", 
                                 bg=self.colors["light"]["card"], fg=self.colors["light"]["subtext"],
                                 font=self.base_font, relief="groove", bd=1)
        api_frame.pack(fill=tk.X, padx=PAD_X, pady=(20, 0), ipady=5)
        
        engine_box = tk.Frame(api_frame, bg=self.colors["light"]["card"])
        engine_box.pack(fill=tk.X, padx=10, pady=(10, 5))
        
        tk.Label(engine_box, text="模型选择:", bg=self.colors["light"]["card"], 
                fg=self.colors["light"]["text"], font=self.base_font).pack(side=tk.LEFT)
        
        self.api_combobox = ttk.Combobox(engine_box, textvariable=self.api_var, values=list(self.api_configs.keys()), 
                    state="readonly", font=self.base_font, width=25)
        self.api_combobox.pack(side=tk.LEFT, padx=10)
        
        self.desc_label = tk.Label(api_frame, text="", bg=self.colors["light"]["card"], 
                                  fg=self.colors["light"]["subtext"], font=self.small_font)
        self.desc_label.pack(anchor="w", padx=10, pady=(0, 10))
        
        def update_engine_desc(*args):
            api = self.api_var.get()
            if api in self.api_configs:
                self.desc_label.config(text=f"说明: {self.api_configs[api]['description']}")
        self.api_var.trace_add("write", update_engine_desc)
        update_engine_desc()
        
        tk.Frame(api_frame, bg=self.colors["light"]["border"], height=1).pack(fill=tk.X, padx=10, pady=5)

        token_box = tk.Frame(api_frame, bg=self.colors["light"]["card"])
        token_box.pack(fill=tk.X, padx=10, pady=(10, 10))
        
        tk.Label(token_box, text="API Token:", bg=self.colors["light"]["card"], 
                fg=self.colors["light"]["text"], font=self.base_font).pack(side=tk.LEFT)
        
        self.token_label = tk.Label(token_box, text="•" * 20 if self.token else "未配置", 
                                   bg=self.colors["light"]["input_bg"], fg=self.colors["light"]["subtext"],
                                   width=25, anchor="w", padx=5, relief="flat")
        self.token_label.pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)
        
        ModernButton(token_box, text="修改 / 配置", command=self.update_token, 
                    variant="secondary").pack(side=tk.RIGHT)

        info_frame = tk.LabelFrame(scrollable_frame, text=" 📖 使用小贴士 ", 
                                  bg=self.colors["light"]["card"], fg=self.colors["light"]["subtext"],
                                  font=self.base_font, relief="groove", bd=1)
        info_frame.pack(fill=tk.X, padx=PAD_X, pady=(20, 20), ipady=5)
        
        tips_container = tk.Frame(info_frame, bg=self.colors["light"]["card"])
        tips_container.pack(fill=tk.X, padx=10, pady=10)
        
        tips_data = [
            ("1.引擎选择", "首选[PaddleOCR-VL]，对双栏/缩进结构还原最强。"),
            ("2.乱序解决", "如Word题目顺序错乱，务必在[识别]页勾选'自动排序'。"),
            ("3.批注技巧", "红笔批注识别受光线影响。建议光线充足，字迹勿太细。"),
            ("4.网络问题", "遇'503'或'网络断开'会自动重试，请耐心等待。"),
            ("5.结果编辑", "Word已应用样式，修改'正文'样式可一键调整格式。")
        ]
        
        self.tip_labels = []
        
        for idx, (title, content) in enumerate(tips_data):
            lbl_title = tk.Label(tips_container, text=title + "：", 
                                bg=self.colors["light"]["card"], fg=self.colors["light"]["text"],
                                font=self.small_font, justify="left", anchor="nw")
            lbl_title.grid(row=idx, column=0, sticky="nw", pady=3, padx=(0, 5))
            
            lbl_content = tk.Label(tips_container, text=content, 
                                  bg=self.colors["light"]["card"], fg=self.colors["light"]["subtext"],
                                  font=self.small_font, justify="left", anchor="nw")
            lbl_content.grid(row=idx, column=1, sticky="nw", pady=3)
            self.tip_labels.append(lbl_content)
            
        tips_container.grid_columnconfigure(1, weight=1) 

        def resize_tips(event):
            target_width = event.width - 140 
            if target_width > 100:
                for lbl in self.tip_labels:
                    lbl.config(wraplength=target_width)
                
        info_frame.bind("<Configure>", resize_tips)

        parent.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

    def setup_recognition_tab(self, parent):
        # 获取当前主题颜色
        colors = self.colors[self.theme_mode]
        parent.configure(bg=colors["card"])
        
        canvas = tk.Canvas(parent, bg=colors["card"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(parent, orient=tk.VERTICAL, command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=colors["card"])
        
        frame_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def _configure_width(event):
            canvas.itemconfig(frame_id, width=event.width)
            
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", _configure_width)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        PAD_X = int(15 * self.scale_factor)
        PAD_Y = int(20 * self.scale_factor)
        
        q_frame = tk.LabelFrame(scrollable_frame, text=" 🔢 智能排序设置 ", 
                               bg=colors["card"], fg=colors["subtext"],
                               font=self.base_font, relief="groove", bd=1)
        q_frame.pack(fill=tk.X, padx=PAD_X, pady=(PAD_Y, 0), ipady=5)
        
        switch_frame = tk.Frame(q_frame, bg=colors["card"])
        switch_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ModernCheckbutton(switch_frame, text="启用题号识别", 
                         variable=self.enable_question_detection).pack(side=tk.LEFT, padx=(0, 20))
        
        ModernCheckbutton(switch_frame, text="自动按题号排序输出", 
                                  variable=self.sort_by_question_var).pack(side=tk.LEFT)
        
        tk.Label(q_frame, text="支持格式:", bg=colors["card"], 
                fg=colors["subtext"], font=self.small_font).pack(anchor="w", padx=10, pady=(5, 0))
        
        fmt_frame = tk.Frame(q_frame, bg=colors["card"])
        fmt_frame.pack(anchor="w", padx=10, pady=5)
        
        formats = [("1.", "1."), ("1、", "1、"), ("(1)", "(1)"), ("一、", "一、"), ("①", "①②③")]
        for text, key in formats:
            if key in self.question_number_formats:
                ModernCheckbutton(fmt_frame, text=text, 
                                variable=self.question_number_formats[key]).pack(side=tk.LEFT, padx=(0, 15))

        struc_frame = tk.LabelFrame(scrollable_frame, text=" 📄 内容与排版 ", 
                                   bg=colors["card"], fg=colors["subtext"],
                                   font=self.base_font, relief="groove", bd=1)
        struc_frame.pack(fill=tk.X, padx=PAD_X, pady=(20, 0), ipady=5)
        
        ModernCheckbutton(struc_frame, text="智能识别选项结构 (自动对齐 A. B. C. D.)", 
                         variable=self.enable_option_detection).pack(anchor="w", padx=10, pady=5)
        
        mode_frame = tk.Frame(struc_frame, bg=colors["card"])
        mode_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(mode_frame, text="排版风格:", bg=colors["card"],
                fg=colors["text"], font=self.base_font).pack(side=tk.LEFT)
        
        self.format_mode_combobox = ttk.Combobox(mode_frame, textvariable=self.output_format_mode, 
                    values=["自动分行", "讲义格式", "保留原样"], state="readonly", 
                    width=15, font=self.base_font)
        self.format_mode_combobox.pack(side=tk.LEFT, padx=10)
        
        self.mode_desc_label = tk.Label(mode_frame, text="", 
                                       bg=colors["card"], fg=colors["subtext"], 
                                       font=self.small_font)
        self.mode_desc_label.pack(side=tk.LEFT)

        ann_frame = tk.LabelFrame(scrollable_frame, text=" 🖊️ 批注提取 (支持多色) ", 
                                 bg=colors["card"], fg=colors["subtext"],
                                 font=self.base_font, relief="groove", bd=1)
        ann_frame.pack(fill=tk.X, padx=PAD_X, pady=(20, 0), ipady=5)
        
        ModernCheckbutton(ann_frame, text="启用颜色检测", 
                         variable=self.enable_annotation_detection).pack(anchor="w", padx=10, pady=5)
        
        colors_frame = tk.Frame(ann_frame, bg=colors["card"])
        colors_frame.pack(fill=tk.X, padx=10, pady=5)
        
        full_color_map = {
            "红色": "#ef4444", 
            "蓝色": "#3b82f6", 
            "绿色": "#22c55e",
            "紫色": "#a855f7",
            "黄色": "#eab308",
            "任意颜色": "#f97316"
        }
        
        col_idx = 0
        row_idx = 0
        for name, hex_val in full_color_map.items():
            if name not in self.selected_colors: continue
            
            cf = tk.Frame(colors_frame, bg=colors["card"])
            cf.grid(row=row_idx, column=col_idx, sticky="w", padx=(0, 15), pady=2)
            
            cv = tk.Canvas(cf, width=12, height=12, highlightthickness=0, bg=colors["card"])
            cv.pack(side=tk.LEFT, padx=(0, 5))
            cv.create_oval(1, 1, 11, 11, fill=hex_val, outline="")
            
            ModernCheckbutton(cf, text=name, variable=self.selected_colors[name]).pack(side=tk.LEFT)
            
            col_idx += 1
            if col_idx > 3:
                col_idx = 0
                row_idx += 1

        filter_frame = tk.LabelFrame(scrollable_frame, text=" 🚫 文本清洗 ", 
                                    bg=colors["card"], fg=colors["subtext"],
                                    font=self.base_font, relief="groove", bd=1)
        filter_frame.pack(fill=tk.X, padx=PAD_X, pady=(20, 20), ipady=5)
        
        tk.Label(filter_frame, text="包含以下关键词的行将被自动删除:", 
                bg=colors["card"], fg=colors["subtext"], 
                font=self.small_font).pack(anchor="w", padx=10, pady=(5, 5))

        # === [核心修复] 使用 Text + ttk.Scrollbar 替换 ScrolledText ===
        text_container = tk.Frame(filter_frame, bg=colors["card"])
        text_container.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        # 1. 创建 ttk 滚动条 (支持深色模式)
        text_scrollbar = ttk.Scrollbar(text_container, orient=tk.VERTICAL)
        
        # 2. 创建标准 Text 控件
        self.blacklist_text = tk.Text(
            text_container, height=3, font=self.base_font,
            bg=colors["input_bg"], fg=colors["text"], 
            relief="flat", bd=1,
            yscrollcommand=text_scrollbar.set
        )
        
        # 3. 绑定
        text_scrollbar.config(command=self.blacklist_text.yview)
        
        # 4. 布局
        text_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.blacklist_text.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.refresh_all_widgets()
        parent.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

    def setup_output_tab(self, parent):
        parent.configure(bg=self.colors["light"]["card"])
        
        canvas = tk.Canvas(parent, bg=self.colors["light"]["card"], highlightthickness=0)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=self.colors["light"]["card"])
        
        frame_id = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        
        def _configure_width(event):
            canvas.itemconfig(frame_id, width=event.width)
            
        scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", _configure_width)
        
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        PAD_X = int(15 * self.scale_factor)
        PAD_Y = int(20 * self.scale_factor)
        
        path_frame = tk.LabelFrame(scrollable_frame, text=" 📁 文件保存位置 ", 
                                  bg=self.colors["light"]["card"], fg=self.colors["light"]["subtext"],
                                  font=self.base_font, relief="groove", bd=1)
        path_frame.pack(fill=tk.X, padx=PAD_X, pady=(PAD_Y, 0), ipady=5)
        
        tk.Label(path_frame, text="保存目录:", bg=self.colors["light"]["card"], 
                fg=self.colors["light"]["text"], font=self.base_font).pack(anchor="w", padx=10, pady=(10, 5))
        
        dir_line = tk.Frame(path_frame, bg=self.colors["light"]["card"])
        dir_line.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        default_dir = os.path.join(self.get_app_path(), "OCR_Output")
        if not self.output_path_var.get():
             self.output_path_var.set(default_dir)
        
        path_entry = tk.Entry(dir_line, textvariable=self.output_path_var, font=self.base_font,
                             bg=self.colors["light"]["input_bg"], fg=self.colors["light"]["text"],
                             relief="flat", bd=1)
        path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5), ipady=4)
        
        ModernButton(dir_line, text="📂 浏览...", variant="secondary", 
                    command=self.select_output_folder).pack(side=tk.RIGHT)
        
        tk.Label(path_frame, text="文件名前缀 (可选):", bg=self.colors["light"]["card"], 
                fg=self.colors["light"]["text"], font=self.base_font).pack(anchor="w", padx=10, pady=(5, 5))
        
        prefix_entry = tk.Entry(path_frame, textvariable=self.prefix_var, font=self.base_font,
                               bg=self.colors["light"]["input_bg"], fg=self.colors["light"]["text"],
                               relief="flat", bd=1)
        prefix_entry.pack(fill=tk.X, padx=10, pady=(0, 15), ipady=4)

        fmt_frame = tk.LabelFrame(scrollable_frame, text=" ⚙️ 格式控制 ", 
                                 bg=self.colors["light"]["card"], fg=self.colors["light"]["subtext"],
                                 font=self.base_font, relief="groove", bd=1)
        fmt_frame.pack(fill=tk.X, padx=PAD_X, pady=(20, 20), ipady=5)
        
        format_line = tk.Frame(fmt_frame, bg=self.colors["light"]["card"])
        format_line.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(format_line, text="文件格式:", bg=self.colors["light"]["card"], 
                fg=self.colors["light"]["text"], font=self.base_font).pack(side=tk.LEFT)
        
        format_options = list(self.output_formats.keys())
        self.format_combobox = ttk.Combobox(format_line, textvariable=self.format_var, values=format_options,
                    state="readonly", font=self.base_font, width=20)
        self.format_combobox.pack(side=tk.LEFT, padx=10)
        
        check_frame = tk.Frame(fmt_frame, bg=self.colors["light"]["card"])
        check_frame.pack(fill=tk.X, padx=10, pady=(5, 15))
        
        ModernCheckbutton(check_frame, text="将所有结果合并为一个文档", 
                                   variable=self.merge_var).pack(anchor="w", pady=5)
        
        ModernCheckbutton(check_frame, text="将检测到的批注文字写入文档", 
                                 variable=self.write_annotations_to_doc).pack(anchor="w", pady=5)
        
      
        parent.bind_all("<MouseWheel>", lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

    def setup_log_tab(self, parent):
        parent.configure(bg=self.colors["light"]["card"])
        
        log_container = tk.Frame(parent, bg=self.colors["light"]["card"])
        log_container.pack(fill=tk.BOTH, expand=True, padx=int(5 * self.scale_factor), pady=int(5 * self.scale_factor))
        
        self.log_text = scrolledtext.ScrolledText(
            log_container,
            height=15,
            font=("Consolas", int(self.base_font_size * self.scale_factor)),
            relief="flat",
            wrap=tk.WORD,
            bg=self.colors["light"]["input_bg"],
            fg=self.colors["light"]["log_text"],
            insertbackground=self.colors["light"]["text"],
            padx=int(8 * self.scale_factor),
            pady=int(8 * self.scale_factor)
        )
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=int(8 * self.scale_factor), pady=int(8 * self.scale_factor))
        
        self.log_text.tag_config("成功", foreground="#22c55e", font=("Consolas", int(self.base_font_size * self.scale_factor)))
        self.log_text.tag_config("错误", foreground="#ef4444", font=("Consolas", int(self.base_font_size * self.scale_factor)))
        self.log_text.tag_config("警告", foreground="#facc15", font=("Consolas", int(self.base_font_size * self.scale_factor)))
        self.log_text.tag_config("信息", foreground=self.colors["light"]["log_text"], font=("Consolas", int(self.base_font_size * self.scale_factor)))
        
        log_control = tk.Frame(log_container, bg=self.colors["light"]["card"])
        log_control.pack(fill=tk.X, padx=int(8 * self.scale_factor), pady=(0, int(8 * self.scale_factor)))
        
        clear_log_btn = ModernButton(log_control, text="清空日志", variant="secondary",
                                    command=self.clear_log, 
                                    padx=int(4 * self.scale_factor),
                                    pady=int(2 * self.scale_factor),
                                    font=self.base_font)
        clear_log_btn.pack(side=tk.RIGHT)
        Tooltip(clear_log_btn, "清空所有日志记录")

    def set_theme(self, mode):
        if self.theme_mode == mode:
            return
        self.theme_mode = mode
        
        zh_mode = self.theme_names.get(mode, mode)
        self.log_message(f"🎨 正在切换主题模式: {zh_mode}", "信息")
        
        if mode == "light":
            self.light_btn.variant = "primary"
            self.dark_btn.variant = "secondary"
        else:
            self.light_btn.variant = "secondary"
            self.dark_btn.variant = "primary"
            
        self.light_btn.update_theme()
        self.dark_btn.update_theme()
        
        self.switch_theme()
        
        self.log_message(f"✅ 已切换到{zh_mode}", "成功")
    
    def change_window_title_bar_color(self, mode):
        try:
            import ctypes
            from ctypes import windll, c_int, byref
            
            if os.name != 'nt':
                return
            
            hwnd = windll.user32.GetParent(self.root.winfo_id())
            
            value = 1 if mode == "dark" else 0
            
            try:
                windll.dwmapi.DwmSetWindowAttribute(hwnd, 20, byref(c_int(value)), 4)
            except:
                try:
                    windll.dwmapi.DwmSetWindowAttribute(hwnd, 19, byref(c_int(value)), 4)
                except:
                    pass
                    
            self.root.update()
        except Exception as e:
            pass

    def log_message(self, msg, level="信息"):
        timestamp = datetime.now().strftime('%H:%M:%S')
        formatted_msg = f"[{timestamp}] {msg}\n"
        
        if not hasattr(self, 'log_text'):
            print(formatted_msg.strip()) 
            return

        if level == "信息":
            self.log_text.insert(tk.END, formatted_msg, "信息")
        elif level == "成功":
            self.log_text.insert(tk.END, formatted_msg, "成功")
        elif level == "错误":
            self.log_text.insert(tk.END, formatted_msg, "错误")
        elif level == "警告":
            self.log_text.insert(tk.END, formatted_msg, "警告")
        else:
            self.log_text.insert(tk.END, formatted_msg, "信息")
        
        self.log_text.see(tk.END)
        try:
            self.log_text.update_idletasks()
        except:
            pass

    def switch_theme(self):
        colors = self.colors[self.theme_mode]
        is_dark = (self.theme_mode == "dark")
        
        style = ttk.Style()
        style.theme_use('clam') 
        
        # Combobox 样式适配
        style.configure("TCombobox",
                       fieldbackground=colors["input_bg"],
                       background=colors["input_bg"],
                       foreground=colors["text"],
                       bordercolor=colors["border"],
                       arrowcolor=colors["text"],
                       relief="flat",
                       borderwidth=1)
        
        style.map("TCombobox",
                  fieldbackground=[("readonly", colors["input_bg"]), 
                                   ("disabled", colors["bg"])],
                  foreground=[("readonly", colors["text"]), 
                              ("disabled", colors["subtext"])],
                  background=[("readonly", colors["input_bg"]), 
                              ("disabled", colors["bg"])],
                  arrowcolor=[("readonly", colors["text"]), 
                              ("disabled", colors["subtext"])])

        self.root.option_add("*TCombobox*Listbox.background", colors["input_bg"])
        self.root.option_add("*TCombobox*Listbox.foreground", colors["text"])
        self.root.option_add("*TCombobox*Listbox.selectBackground", colors["primary"])
        self.root.option_add("*TCombobox*Listbox.selectForeground", "#ffffff")

        scrollbar_bg = "#333333" if is_dark else colors["subtext"]
        style.configure("Vertical.TScrollbar", background=scrollbar_bg, troughcolor=colors["bg"], 
                       bordercolor=colors["bg"], arrowcolor=colors["text"], relief="flat")
        style.map("Vertical.TScrollbar", background=[("active", colors["primary"])])

        # === [核心修复] Notebook 完美去白边 ===
        # 关键：将 lightcolor 和 darkcolor 强制设为背景色
        # 这样 clam 主题绘制的 3D 边框就会"隐形"
        style.configure("TNotebook", 
                       background=colors["bg"], 
                       borderwidth=0,
                       lightcolor=colors["bg"], # 关键！消除左侧和上侧白线
                       darkcolor=colors["bg"])  # 关键！消除右侧和下侧白线

        style.configure("TNotebook.Tab", background=colors["tab_bg"], foreground=colors["tab_inactive"], padding=[8, 4])
        style.map("TNotebook.Tab", background=[("selected", colors["bg"])], foreground=[("selected", colors["tab_active"])])
        
        style.configure("TFrame", background=colors["bg"])
        
        def update_widget_tree(widget, force_bg=None):
            try:
                if isinstance(widget, (ModernButton, ModernCheckbutton)):
                    return
                    
                if isinstance(widget, (ttk.Combobox, ttk.Progressbar, ttk.Notebook, ttk.Scrollbar, ttk.PanedWindow)):
                    return

                if hasattr(widget, 'configure'):
                    if force_bg:
                        if isinstance(widget, (tk.Text, scrolledtext.ScrolledText, tk.Entry)):
                            widget.configure(bg=colors["input_bg"])
                        elif isinstance(widget, (tk.Canvas, tk.Frame, tk.LabelFrame, tk.Label)):
                             widget.configure(bg=force_bg)
                    else:
                        if isinstance(widget, (tk.Frame, tk.LabelFrame)):
                             try:
                                 parent_bg = widget.master.cget("bg") if widget.master else colors["bg"]
                                 widget.configure(bg=parent_bg)
                             except:
                                 widget.configure(bg=colors["bg"])

                if hasattr(widget, "configure"):
                    try:
                        # 修复 Highlight 边框颜色
                        if int(str(widget.cget("highlightthickness"))) > 0:
                            widget.configure(highlightbackground=colors["border"])
                    except: pass

                if hasattr(widget, 'configure') and hasattr(widget, 'cget'):
                    try:
                        if isinstance(widget, (tk.Label, tk.Entry, tk.Text, scrolledtext.ScrolledText)):
                             widget.configure(fg=colors["text"])
                        
                        if 'insertbackground' in widget.keys():
                            widget.configure(insertbackground=colors["text"])
                    except: pass
            except: pass
            
            for child in widget.winfo_children():
                update_widget_tree(child, force_bg=force_bg)

        self.root.configure(bg=colors["bg"])
        
        if hasattr(self, 'main_container'):
            self.main_container.configure(bg=colors["bg"])
            for w in self.main_container.winfo_children():
                update_widget_tree(w, force_bg=None)
        
        card_areas = ['left_panel', 'control_bar']
        if hasattr(self, 'notebook'):
            for tab in self.notebook.winfo_children():
                # 确保 tab 内容页也是 Card 色
                tab.configure(bg=colors["card"])
                for child in tab.winfo_children():
                    update_widget_tree(child, force_bg=colors["card"])
        
        for area_name in card_areas:
            if hasattr(self, area_name):
                area = getattr(self, area_name)
                area.configure(bg=colors["card"], highlightbackground=colors["border"])
                for child in area.winfo_children():
                    update_widget_tree(child, force_bg=colors["card"])

        if hasattr(self, 'log_text'):
             self.log_text.configure(bg=colors["input_bg"], fg=colors["log_text"], insertbackground=colors["text"])
        
        self.update_all_comboboxes()
        
        self.refresh_all_widgets()
        self.root.update_idletasks()
        
        self.change_window_title_bar_color(self.theme_mode)

    def update_all_comboboxes(self):
        def update_combobox(widget):
            if isinstance(widget, ttk.Combobox):
                widget.configure(style="TCombobox")
            for child in widget.winfo_children():
                update_combobox(child)
        
        update_combobox(self.root)

    def refresh_all_widgets(self):
        for widget in self.refreshables:
            try:
                if hasattr(widget, '_draw'):
                    widget._draw()
                elif hasattr(widget, 'update_theme'):
                    widget.update_theme()
            except Exception:
                pass

    def get_selected_colors(self):
        if not self.enable_annotation_detection.get():
            return []
        
        return [color_name for color_name, var in self.selected_colors.items() 
                if var.get() and color_name in ["红色", "蓝色", "绿色", "紫色", "黄色", "任意颜色"]]

    def preview_annotation_detection(self):
        if not self.files:
            messagebox.showwarning("提示", "请先添加文件")
            return
        
        if not self.has_opencv:
            messagebox.showerror("错误", "OpenCV未安装，无法使用批注检测功能")
            return
        
        if not self.enable_annotation_detection.get():
            messagebox.showwarning("提示", "请先在'识别设置'中启用批注检测功能")
            return
        
        file_path = self.files[0]
        file_name = os.path.basename(file_path)
        
        if not os.path.exists(file_path):
            self.log_message(f"❌ 文件不存在: {file_path}", "错误")
            return
        
        self.log_message(f"🔍 正在检测批注: {file_name}", "信息")
        
        try:
            selected_colors = self.get_selected_colors()
            if not selected_colors:
                messagebox.showwarning("提示", "请至少选择一种批注颜色")
                return
            
            tolerance = self.color_tolerance.get()
            threshold = 0.7 - (tolerance / 100) * 0.2
            
            result = AnnotationDetector.detect_annotations(
                file_path, 
                selected_colors=selected_colors,
                threshold=max(0.5, threshold)
            )
            
            if result is not None and result['annotations']:
                highlighted_path = AnnotationDetector.highlight_annotations(
                    file_path, 
                    selected_colors=selected_colors
                )
                
                color_counts = {}
                for ann in result['annotations']:
                    color = ann['color']
                    color_counts[color] = color_counts.get(color, 0) + 1
                
                total_count = len(result['annotations'])
                color_info = ", ".join([f"{color}{count}处" for color, count in color_counts.items()])
                
                self.log_message(f"✅ 检测到批注{total_count}处 ({color_info})", "成功")
                
                if os.path.exists(highlighted_path):
                    if os.name == 'nt':
                        os.startfile(highlighted_path)
                    elif os.name == 'posix':
                        subprocess.run(['open', highlighted_path] if sys.platform == 'darwin' else ['xdg-open', highlighted_path])
                
                messagebox.showinfo("批注检测结果", 
                                  f"检测完成！\n\n文件: {file_name}\n检测到批注总数: {total_count}\n\n按颜色统计:\n{color_info}\n\n已保存高亮图像到: {highlighted_path}")
            else:
                self.log_message("ℹ️ 未检测到批注", "信息")
                messagebox.showinfo("批注检测结果", "未检测到批注")
                
        except Exception as e:
            self.log_message(f"❌ 批注检测失败: {str(e)}", "错误")
            messagebox.showerror("错误", f"批注检测失败:\n{str(e)}")

    def start_processing(self):
        self.disable_tooltips = True
        
        if not self.token:
            self.disable_tooltips = False
            messagebox.showerror("错误", "未配置 API Token！\n请点击'配置'选项卡设置 Token。")
            return
            
        if not self.files:
            self.disable_tooltips = False
            messagebox.showwarning("提示", "请先添加文件")
            return
        
        self.processing = True 
        
        self.output_prefix = self.prefix_var.get().strip()
        
        self.start_button.config(state="disabled", text="⏳ 处理中...")
        self.stop_button.config(state="normal")
        self.show_progress_bar()
        
        threading.Thread(target=self.process_thread, daemon=True).start()

    def show_progress_bar(self):
        self.progress_container.config(height=int(30 * self.scale_factor))
        
        for widget in self.progress_container.winfo_children():
            widget.destroy()
        
        progress_frame = tk.Frame(self.progress_container, bg=self.colors[self.theme_mode]["bg"])
        progress_frame.pack(fill=tk.BOTH, expand=True, padx=int(10 * self.scale_factor), pady=int(5 * self.scale_factor))
        
        progress_label = tk.Label(
            progress_frame,
            text="⏳ 正在处理...",
            bg=self.colors[self.theme_mode]["bg"],
            fg=self.colors[self.theme_mode]["warning"],
            font=self.base_font
        )
        progress_label.pack(side=tk.LEFT)
        
        self.progress_bar = ttk.Progressbar(
            progress_frame, 
            mode="indeterminate",
            length=int(260 * self.scale_factor)
        )
        self.progress_bar.pack(side=tk.RIGHT, fill=tk.X, expand=True, padx=(int(10 * self.scale_factor), 0))
        self.progress_bar.start(10)

    def update_preview_display(self):
       
        for widget in self.preview_content.winfo_children():
            widget.destroy()
        
        c = self.colors[self.theme_mode]
        row, col = 0, 0
        max_cols = 4
        
        for f in self.files:
            preview_frame = tk.Frame(
                self.preview_content,
                bg=c["card"],
                highlightthickness=1,
                highlightbackground=c["border"],
                relief="solid"
            )
            preview_frame.grid(row=row, column=col, padx=int(5 * self.scale_factor), 
                               pady=int(5 * self.scale_factor), sticky="nsew")
            
            self.preview_content.grid_columnconfigure(col, weight=1)
            
            if f in self.file_previews:
                img_label = tk.Label(preview_frame, image=self.file_previews[f], bg=c["card"])
                img_label.pack(pady=(int(5 * self.scale_factor), int(2 * self.scale_factor)))
            
            filename = os.path.basename(f)
            if len(filename) > 12:
                filename = filename[:10] + "..."
            
            name_label = tk.Label(
                preview_frame,
                text=filename,
                bg=c["card"],
                fg=c["text"],
                font=self.small_font
            )
            name_label.pack(pady=(0, int(2 * self.scale_factor)))
            
            del_btn = tk.Button(
                preview_frame,
                text="× 删除",
                bg="#fee2e2",
                fg="#ef4444",
                relief="flat",
                bd=0,
                font=(self.font_family, int(7 * self.scale_factor)),
                cursor="hand2",
                command=lambda p=f: self.delete_single_file(p)
            )
            del_btn.pack(pady=(0, int(5 * self.scale_factor)), ipadx=5)
            
            col += 1
            if col >= max_cols:
                col = 0
                row += 1
        
        self.preview_content.update_idletasks()
        self.preview_canvas.configure(scrollregion=self.preview_canvas.bbox("all"))

    def delete_single_file(self, file_path):
        if file_path in self.files:
            self.files.remove(file_path)
            if file_path in self.file_previews:
                del self.file_previews[file_path]
            self.update_preview_display()
            self.log_message(f"🗑️已移除: {os.path.basename(file_path)}", "信息")
    
    def clear_list(self):
        self.files = []
        self.file_previews = {}
        self.update_preview_display()
        self.log_message("🗑️文件列表已清空", "信息")

    def hide_progress_bar(self):
        try:
            if hasattr(self, 'progress_bar') and self.progress_bar:
                self.progress_bar.stop()
        except Exception:
            pass 
            
        try:
            if hasattr(self, 'progress_container') and self.progress_container:
                self.progress_container.config(height=0)
                for widget in self.progress_container.winfo_children():
                    widget.destroy()
        except Exception:
            pass

    def process_thread(self):
        out_dir = self.output_path_var.get()
        if not os.path.exists(out_dir):
            os.makedirs(out_dir)
        
        results = []
        total = len(self.files)
        success_count = 0
        current_model = self.api_var.get()
        
        for i, fpath in enumerate(self.files):
            if not self.processing: break 
            
            fname = os.path.basename(fpath)
            self.log_message(f"🚀 [第 {i+1}/{total} 个] 开始处理: {fname}", "信息")
            self.root.after(0, lambda f=fname, idx=i+1: self.update_progress_label(f"正在处理 ({idx}/{total}): {f}"))
            
            try:
                if not self.processing: break
                
                api_result = self.call_ocr_api_with_retry(fpath, retries=3, timeout=(60, 300))
                
                if not self.processing: break
                if not api_result:
                    self.log_message(f"❌ 识别失败: {fname}", "错误")
                    continue
                
                text_blocks = self.extract_text_blocks(api_result) 
                processed_content = self.intelligent_processing(api_result)
                
                annotations_info = None
                if self.has_opencv and self.enable_annotation_detection.get():
                    if not self.processing: break
                    
                    selected_colors = self.get_selected_colors()
                    if selected_colors:
                        result = AnnotationDetector.detect_annotations(fpath, selected_colors)
                        if result and result['annotations']:
                             pass

                if not self.processing: break

                processed_content = f"【📄 来源文件: {fname}】\n" + processed_content

                if self.sort_by_question_var.get():
                    processed_content = self.sort_content_by_question(processed_content)
                
                results.append({
                    'file_path': fpath,
                    'content': processed_content,
                    'text_blocks': text_blocks,
                    'annotations': annotations_info
                })
                success_count += 1
                self.log_message(f"💾 处理就绪: {fname}", "成功")
                
            except Exception as e:
                self.log_message(f"⚠️ 处理异常 {fname}: {str(e)}", "错误")
        
        if self.processing and results:
            try:
                self.log_message(f"📂 正在保存结果...", "信息")
                self.save_results(results, out_dir)
            except Exception as e:
                self.log_message(f"❌ 保存失败: {str(e)}", "错误")
        
        was_stopped = not self.processing
        self.processing = False
        self.root.after(0, lambda: self.reset_ui(success_count, total, was_stopped))

    def update_progress_label(self, text):
        for widget in self.progress_container.winfo_children():
            if isinstance(widget, tk.Frame):
                for child in widget.winfo_children():
                    if isinstance(child, tk.Label) and "正在处理" in child.cget("text"):
                        child.config(text=text)
                        break

    def clean_v5_text(self, text):
        if not text: return ""
        import re
        replacements = [
            ('改柄句', ''), ('搭配不当', ''), ('前后矛盾', ''), ('搭配不', ''),
            ('\u3000', ' ')
        ]
        result = str(text)
        for old, new in replacements:
            result = result.replace(old, new)
        
        result = result.replace('\n', ' ').replace('\r', ' ')
        return re.sub(r'\s+', ' ', result).strip()

    def extract_text_blocks(self, api_result):
        """
        [终极修复版] 针对 VL 模型 prunedResult 结构进行穿透处理
        """
        text_blocks = []
        try:
            # 1. 调试保存 (保持不变)
            import json
            with open("debug_last_response.json", "w", encoding="utf-8") as f:
                json.dump(api_result, f, ensure_ascii=False, indent=2)
            
            data = api_result
            if isinstance(data, str):
                try: data = json.loads(data)
                except: pass
            
            if isinstance(data, dict):
                if 'result' in data: data = data['result']
                if isinstance(data, dict) and 'result' in data: data = data['result']

            # ---------------------------------------------------------
            # 场景 A: 针对 V5 模型 (ocrResults) - 你之前的代码这里是对的
            # ---------------------------------------------------------
            if isinstance(data, dict) and 'ocrResults' in data:
                for item in data['ocrResults']:
                    target_obj = item.get('prunedResult', item) # 兼容 V5
                    if 'rec_texts' in target_obj and 'rec_boxes' in target_obj:
                        texts = target_obj['rec_texts']
                        boxes = target_obj['rec_boxes']
                        for i in range(min(len(texts), len(boxes))):
                            text_blocks.append({
                                'bbox': self._normalize_bbox(boxes[i]),
                                'text': self.clean_v5_text(str(texts[i])),
                                'source': 'V5-Standard'
                            })

            # ---------------------------------------------------------
            # 场景 B: 针对 VL 模型 (layoutParsingResults) - 【这里是之前的 Bug 所在】
            # ---------------------------------------------------------
            targets = []
            if isinstance(data, dict) and 'layoutParsingResults' in data: 
                targets = data['layoutParsingResults']
            
            if isinstance(targets, list) and targets:
                for res in targets:
                    # === [核心修复] 增加对 prunedResult 的剥离 ===
                    # 你的 debug.json 显示 parsing_res_list 在 prunedResult 里面
                    actual_res = res.get('prunedResult', res) 
                    
                    parsing_list = actual_res.get('parsing_res_list', [])
                    for item in parsing_list:
                        text = item.get('block_content') or item.get('text')
                        bbox = item.get('block_bbox') or item.get('bbox')
                        if text and bbox:
                            text_blocks.append({
                                'bbox': self._normalize_bbox(bbox),
                                'text': str(text),
                                'source': 'VL-Standard' # 标记为标准解析成功
                            })

            # ---------------------------------------------------------
            # 兜底: 深度搜索 (只有当上面都失败时才运行)
            # ---------------------------------------------------------
            if not text_blocks:
                self.log_message("⚠️ 标准解析未命中，启用深度搜索...", "警告")
                text_blocks = self._deep_search_blocks(data)

            # 补充中心点数据用于排序
            for b in text_blocks:
                x, y, w, h = b['bbox']
                b['center'] = (x + w//2, y + h//2)
                b['y_range'] = (y, y + h)

        except Exception as e:
            self.log_message(f"⚠️ 提取过程异常: {e}", "错误")
        
        return text_blocks

    def _normalize_bbox(self, box):
        if not box or len(box) < 4: return [0,0,0,0]
        v1, v2, v3, v4 = map(int, box[:4])
        # 兼容 [x1, y1, x2, y2] 格式
        if v3 > v1 and v4 > v2: return [v1, v2, v3-v1, v4-v2]
        return [v1, v2, v3, v4]

    def _deep_search_blocks(self, data):
        found = []
        if isinstance(data, dict):
            if 'text' in data and 'bbox' in data:
                found.append({'bbox': self._normalize_bbox(data['bbox']), 'text': str(data['text'])})
            elif 'block_content' in data and 'block_bbox' in data:
                found.append({'bbox': self._normalize_bbox(data['block_bbox']), 'text': str(data['block_content'])})
            for v in data.values(): found.extend(self._deep_search_blocks(v))
        elif isinstance(data, list):
            for v in data: found.extend(self._deep_search_blocks(v))
        return found

    def intelligent_processing(self, result):
        try:
            text_blocks = self.extract_text_blocks(result)
            if not text_blocks: return ""

            blacklist = []
            if hasattr(self, 'blacklist_text'):
                raw_bl = self.blacklist_text.get("1.0", tk.END)
                blacklist = [line.strip() for line in raw_bl.split('\n') if line.strip()]

            text_blocks.sort(key=lambda b: b['bbox'][1])
            sorted_lines = []
            current_line = [text_blocks[0]]
            
            for i in range(1, len(text_blocks)):
                b = text_blocks[i]
                last_b = current_line[-1]
                
                y_diff = abs(b['bbox'][1] - last_b['bbox'][1])
                height_avg = (b['bbox'][3] + last_b['bbox'][3]) / 2
                
                if y_diff < (height_avg * 0.5): 
                    current_line.append(b)
                else:
                    current_line.sort(key=lambda x: x['bbox'][0])
                    sorted_lines.extend(current_line)
                    current_line = [b]
            
            if current_line:
                current_line.sort(key=lambda x: x['bbox'][0])
                sorted_lines.extend(current_line)

            final_lines = []
            for b in sorted_lines:
                text = b['text'].strip()
                if not text: continue
                
                is_spam = False
                for block_word in blacklist:
                    if block_word in text:
                        is_spam = True
                        break
                if is_spam:
                    continue
                
                final_lines.append(text)

            return "\n".join(final_lines)

        except Exception as e:
            self.log_message(f"解析内容出错: {e}", "错误")
            return ""

    def match_annotations_to_text_blocks(self, text_blocks, annotations):
        if not text_blocks or not annotations: return {}
        
        matches = {}
        
        for ann in annotations:
            ax, ay, aw, ah = ann['bbox']
            a_center_y = ay + ah // 2
            a_area = aw * ah
            
            best_match = None
            best_score = -1
            
            for block_idx, block in enumerate(text_blocks):
                bx, by, bw, bh = block['bbox']
                
                overlap_x1 = max(ax, bx)
                overlap_y1 = max(ay, by)
                overlap_x2 = min(ax + aw, bx + bw)
                overlap_y2 = min(ay + ah, by + bh)
                
                iou_score = 0
                if overlap_x2 > overlap_x1 and overlap_y2 > overlap_y1:
                    overlap_area = (overlap_x2 - overlap_x1) * (overlap_y2 - overlap_y1)
                    union_area = a_area + (bw * bh) - overlap_area
                    if union_area > 0:
                        iou_score = overlap_area / union_area
                
                b_center_y = by + bh // 2
                v_dist = abs(a_center_y - b_center_y)
                v_score = max(0, 1 - v_dist / (bh * 2))
                
                if iou_score > 0:
                    total_score = iou_score * 0.7 + v_score * 0.3 + 1.0
                else:
                    if v_dist < bh * 1.5:
                        total_score = v_score
                    else:
                        total_score = 0

                if total_score > best_score:
                    best_score = total_score
                    best_match = block_idx
            
            if best_match is not None and best_score > 0.4:
                if best_match not in matches: matches[best_match] = []
                matches[best_match].append(ann)
                
        return matches

    def sort_content_by_question(self, content):
        try:
            lines = content.split('\n')
            
            question_patterns = [
                r'^(\d+)[\.、．]\s*.*',
                r'^\((\d+)\)\s*.*',
                r'^([一二三四五六七八九十]+)[、.．]\s*.*',
                r'^([①②③④⑤⑥⑦⑧⑨⑩]+)\s*.*',
                r'^第(\d+)题.*'
            ]
            
            blocks = []
            current_block = {"id": 0, "lines": []}
            
            def get_question_id(text):
                for p in question_patterns:
                    match = re.match(p, text.strip())
                    if match:
                        num_str = match.group(1)
                        cn_map = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
                        if num_str in cn_map: return cn_map[num_str]
                        circle_map = {'①':1,'②':2,'③':3,'④':4,'⑤':5,'⑥':6,'⑦':7,'⑧':8,'⑨':9,'⑩':10}
                        if num_str in circle_map: return circle_map[num_str]
                        try: return int(num_str)
                        except: return 9999
                return None

            for line in lines:
                line = line.strip()
                if not line: continue
                
                if line.startswith("【📄 来源文件:"):
                    current_block["lines"].append(line)
                    continue
                    
                qid = get_question_id(line)
                
                if qid is not None:
                    if current_block["lines"]:
                        blocks.append(current_block)
                    current_block = {"id": qid, "lines": [line]}
                else:
                    current_block["lines"].append(line)
            
            if current_block["lines"]:
                blocks.append(current_block)
            
            header_blocks = [b for b in blocks if b["id"] == 0]
            question_blocks = [b for b in blocks if b["id"] != 0]
            
            question_blocks.sort(key=lambda x: x["id"])
            
            final_lines = []
            
            if header_blocks:
                final_lines.append("【⚠️ 未识别到题号的内容 / 卷头信息】")
                for b in header_blocks:
                    final_lines.extend(b["lines"])
                final_lines.append("-" * 30)
            
            for b in question_blocks:
                final_lines.extend(b["lines"])
                final_lines.append("") 
            
            return '\n'.join(final_lines)
            
        except Exception as e:
            self.log_message(f"⚠️ 排序出错: {str(e)}", "警告")
            return content

    def smart_global_sort(self, results):
        """
        [智能全局排序]
        核心逻辑：解析所有文件的所有题目，将"非题号内容"(如文件名、卷头)
        吸附在"下一道题"身上，然后统一按题号排序。
        这样既能全局排序，又不会丢失文件名和卷头。
        """
        try:
            all_blocks = []
            
            # 题号正则
            question_patterns = [
                r'^\s*(\d+)[\.、．\s]',       # "1. "
                r'^\s*\(\s*(\d+)\s*\)',      # "(1)"
                r'^\s*([一二三四五六七八九十]+)[\.、．\s]', # "一、"
                r'^\s*第\s*(\d+)\s*题'       # "第1题"
            ]
            
            def get_question_id(text_line):
                for p in question_patterns:
                    import re
                    match = re.match(p, text_line)
                    if match:
                        num_str = match.group(1)
                        cn_map = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
                        if num_str in cn_map: return cn_map[num_str]
                        try: return int(num_str)
                        except: return 99999
                return None

            # 遍历所有文件的结果
            for res in results:
                content = res['content']
                lines = content.split('\n')
                
                # 临时缓冲区，用来存"还没遇到题号的文字" (比如文件名、卷头)
                pending_headers = []
                
                # 当前正在处理的题目块
                current_block = None
                
                for line in lines:
                    clean_line = line.strip()
                    if not clean_line: continue
                    
                    qid = get_question_id(line)
                    
                    if qid is not None:
                        # === 发现新题目 ===
                        
                        # 1. 先保存上一个题目块
                        if current_block:
                            all_blocks.append(current_block)
                        
                        # 2. 创建新题目块
                        # 重点：把之前积攒的 headers (文件名/卷头) 全部粘在这个新题目头上
                        full_text_lines = pending_headers + [line]
                        current_block = {
                            "id": qid,
                            "lines": full_text_lines,
                            "sort_key": qid
                        }
                        # 清空缓冲区
                        pending_headers = []
                    else:
                        # === 不是题目 (是选项、正文、或者文件名) ===
                        if current_block:
                            # 如果已经在一个题目里了，就追加到题目后面 (作为选项/内容)
                            current_block["lines"].append(line)
                        else:
                            # 如果还没遇到任何题目 (说明是文件最开头的卷头/文件名)
                            # 先攒着，等遇到下一道题时，粘给下一道题
                            pending_headers.append(line)
                
                # 循环结束，保存最后一个块
                if current_block:
                    all_blocks.append(current_block)
                
                # 如果最后还有剩下的 headers (比如文件末尾的页码)，就粘到最后一个块屁股后面
                if pending_headers and all_blocks:
                    all_blocks[-1]["lines"].extend(pending_headers)
                elif pending_headers:
                    # 万一这个文件全是废话，没有一道题，就作为一个 ID=0 的块存起来
                    all_blocks.append({"id": 0, "lines": pending_headers, "sort_key": 0})

            # === 核心：对所有块进行全局排序 ===
            # 按题号从小到大排
            all_blocks.sort(key=lambda x: x["sort_key"])
            
            # === 重组 ===
            final_lines = []
            for b in all_blocks:
                final_lines.extend(b["lines"])
                final_lines.append("") # 题间空行
                
            return '\n'.join(final_lines)

        except Exception as e:
            self.log_message(f"全局排序出错: {e}", "错误")
            # 出错兜底：直接硬拼
            return "\n\n".join([r['content'] for r in results])

    def save_as_docx(self, results, output_path):
        doc = Document()
        styles = doc.styles
        normal_style = styles['Normal']
        normal_style.font.name = 'Microsoft YaHei UI'
        normal_style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei UI')
        
        for i, result in enumerate(results):
            if result['file_path'] != 'MERGED_VIRTUAL_FILE':
                if i > 0: doc.add_page_break()
                filename = os.path.basename(result['file_path'])
                doc.add_heading(filename, level=2)
            
            content = result['content']
            lines = content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line: continue
                
                if line.startswith("【📄 来源文件:"):
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                    run.font.size = Pt(9)
                    run.italic = True
                    p.paragraph_format.space_before = Pt(12)
                    continue

                p = doc.add_paragraph()
                
                if '**' in line:
                    parts = line.split('**')
                    for idx, part in enumerate(parts):
                        run = p.add_run(part)
                        if idx % 2 == 1: 
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    p.add_run(line)
        
        doc.save(output_path)
        self.log_message(f"✅ Word文档生成成功: {os.path.basename(output_path)}", "成功")
        
    def save_results(self, results, out_dir):
        ext = self.output_formats[self.format_var.get()]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        prefix = self.output_prefix + "_" if self.output_prefix else ""
        
        if self.output_format_mode.get() == "讲义格式":
            results = self.format_as_handout(results)
        
        if self.merge_var.get():
            fname = f"{prefix}合并结果_{timestamp}.{ext}"
            output_path = os.path.join(out_dir, fname)
            
            # === [核心修改] 判断是否需要全局排序 ===
            if self.sort_by_question_var.get():
                self.log_message("🔄 执行智能全局重排 (解决文件乱序问题)...", "信息")
                final_content = self.smart_global_sort(results)
            else:
                # 没勾选排序，就按文件列表顺序硬拼
                final_content = "\n\n".join([str(r['content']) for r in results])
            
            # 构造结果对象
            merged_result = [{
                'file_path': 'MERGED_VIRTUAL_FILE',
                'content': final_content,
                'text_blocks': [],
                'annotations': None
            }]

            try:
                if ext == "docx":
                    self.save_as_docx(merged_result, output_path)
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(final_content)
                
                self.log_message(f"✅ 已保存合并文件: {fname}", "成功")
            except Exception as e:
                raise e
        else:
            # 不合并的情况 (保持原样)
            for result in results:
                base = os.path.splitext(os.path.basename(result['file_path']))[0]
                fname = f"{prefix}{base}.{ext}"
                output_path = os.path.join(out_dir, fname)
                try:
                    # 单文件模式下，如果选了排序，建议在 process_thread 里已经排好(用上一轮代码)，这里只管保存
                    if ext == "docx":
                        self.save_as_docx([result], output_path)
                    else:
                        with open(output_path, 'w', encoding='utf-8') as f:
                            f.write(str(result['content']))
                except Exception as e:
                    self.log_message(f"❌ 保存失败 {fname}: {e}", "错误")
    
    def format_as_handout(self, results):
        formatted_results = []
        
        for result in results:
            content = result['content']
            lines = content.split('\n')
            formatted_lines = []
            
            for line in lines:
                line = line.strip()
                if not line: 
                    formatted_lines.append("")
                    continue

                if self.enable_question_detection.get():
                    if re.match(r'^(\d+|[一二三四五]+|\([0-9]+\)|[①-⑩])[\.、．)]', line):
                        line = re.sub(r'^([^\s]+)(.*)', r'**\1** \2', line)
                
                if self.enable_option_detection.get():
                    if re.match(r'^[A-D][\.、．]', line) or re.match(r'^\([A-D]\)', line):
                        line = f"    {line}"
                        line = re.sub(r'^\s+([A-D][\.、．]|\([A-D]\))', r'    **\1**', line)
                
                formatted_lines.append(line)
            
            formatted_content = '\n'.join(formatted_lines)
            formatted_results.append({
                'file_path': result['file_path'],
                'content': formatted_content,
                'text_blocks': result.get('text_blocks', []),
                'annotations': result.get('annotations')
            })
        
        return formatted_results
    
    def select_images(self):
        files = filedialog.askopenfilenames(
            title="选择图片或PDF文件",
            filetypes=[
                ("图片文件", "*.png *.jpg *.jpeg *.bmp"),
                ("PDF文件", "*.pdf"),
                ("所有文件", "*.*")
            ]
        )
        self.add_files(files)
        
        if files and self.has_opencv:
            for file_path in files:
                if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
                    self.auto_update_annotation_colors(file_path)
                    break

    def select_folder(self):
        folder = filedialog.askdirectory(title="选择包含图片/PDF的文件夹")
        if folder:
            files = []
            for root, _, filenames in os.walk(folder):
                for f in filenames:
                    if f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.pdf')):
                        files.append(os.path.join(root, f))
            self.add_files(files)
            
            if files and self.has_opencv:
                for file_path in files:
                    if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
                        self.auto_update_annotation_colors(file_path)
                        break

    def add_files(self, paths):
        for p in paths:
            if p not in self.files:
                self.files.append(p)
                self.create_preview(p)
        
        self.log_message(f"📄 当前列表共 {len(self.files)} 个文件", "信息")
        self.update_preview_display()

    def auto_update_annotation_colors(self, image_path):
        if not self.has_opencv:
            return
        
        def analyze_colors():
            try:
                colors = AnnotationDetector.extract_dominant_annotation_colors(image_path, max_colors=2)
                
                if not colors:
                    return
                
                self.root.after(0, lambda: self._update_color_selection(colors))
                
                self.log_message(
                    f"🎨已自动识别批注颜色: {', '.join(colors)}",
                    "信息"
                )
            except Exception as e:
                pass
        
        threading.Thread(target=analyze_colors, daemon=True).start()

    def _update_color_selection(self, colors):
        self.enable_annotation_detection.set(True)
        
        for color_name, var in self.selected_colors.items():
            if color_name in ["红色", "蓝色", "绿色", "紫色", "黄色", "任意颜色"]:
                var.set(color_name in colors)

    def create_preview(self, path):
        try:
            size = (int(80 * self.scale_factor), int(80 * self.scale_factor))
            if path.lower().endswith('.pdf'):
                img = Image.new('RGB', size, color='#3b82f6')
                draw = ImageDraw.Draw(img)
                try:
                    draw_font = ImageFont.truetype("arial.ttf", 16)
                except:
                    draw_font = ImageFont.load_default()
                draw.text((size[0]//2-15, size[1]//2-8), "PDF", fill="white", font=draw_font)
            else:
                img = Image.open(path)
                img = img.convert("RGB")
                img.thumbnail(size, Image.Resampling.LANCZOS)
            
            self.file_previews[path] = ImageTk.PhotoImage(img)
        except Exception as e:
            pass

    def select_output_folder(self):
        folder = filedialog.askdirectory(title="选择输出文件夹")
        if folder:
            self.output_path_var.set(folder)
            self.log_message(f"📁 输出路径已更新: {folder}", "信息")

    def on_mousewheel(self, event):
        self.preview_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def reset_ui(self, success_count, total, was_stopped=False):
        self.disable_tooltips = False
        self.hide_progress_bar()
        
        self.start_button.config(state="normal", text="▶ 开始处理 (F5)")
        self.stop_button.config(state="disabled")
        
        if was_stopped:
            self.log_message("⚠️ 任务已强制停止", "警告")
        elif success_count > 0:
            self.log_message(f"🎉 全部完成！成功处理 {success_count}/{total} 个文件", "成功")
            messagebox.showinfo("处理完成", f"任务完成！\n成功: {success_count}/{total}\n保存路径: {self.output_path_var.get()}")
        else:
            self.log_message("⚠️ 任务结束，未生成有效文件", "警告")
            messagebox.showwarning("提示", "未成功处理任何文件，请检查日志。")

    def open_output_folder(self):
        path = self.output_path_var.get()
        if os.path.exists(path):
            if os.name == 'nt':
                os.startfile(path)
            elif os.name == 'posix':
                if sys.platform == 'darwin':
                    subprocess.run(['open', path])
                else:
                    subprocess.run(['xdg-open', path])
        else:
            messagebox.showinfo("提示", "输出文件夹不存在，请先选择有效的输出路径")

    def setup_shortcuts(self):
        self.root.bind('<F5>', lambda e: self.start_processing())
        self.root.bind('<Delete>', lambda e: self.clear_list())
        self.root.bind('<Control-o>', lambda e: self.select_images())
        self.root.bind('<Control-O>', lambda e: self.select_folder())

    def stop_processing(self):
        if not self.processing: return
        
        self.processing = False 
        self.log_message("🛑 正在停止任务，请稍候...", "警告")
        
        self.hide_progress_bar()
        self.start_button.config(state="normal", text="▶ 开始处理 (F5)")
        self.stop_button.config(state="disabled")
        self.disable_tooltips = False

    def on_closing(self):
        if self.processing:
            if messagebox.askyesno("确认", "当前有任务正在处理，确定要退出吗？"):
                self.processing = False
                time.sleep(0.5)
                self.root.destroy()
        else:
            self.root.destroy()

    def update_token(self):
        config_file = os.path.join(self.get_app_path(), "config.ini")
        
        if not os.path.exists(config_file):
            result = messagebox.askyesno("配置文件不存在", "是否创建配置文件模板？")
            if result:
                self.create_config_template(config_file, "")
            else:
                return
        
        # 1. 创建弹窗
        dialog = tk.Toplevel(self.root)
        dialog.title("更新API Token")
        dialog.app = self
        
        # === [核心修复1] 强制显示标题栏 Logo ===
        # Windows 特性：transient 窗口会自动隐藏标题栏图标
        # 所以必须注释掉 transient，才能让 Logo 显示出来
        # dialog.transient(self.root)  # <--- 已注释，勿开
        
        # === [核心修复2] 完整的图标加载策略 ===
        # 策略A: 如果主程序加载了 PNG 图标 (self.app_icon)，直接复用
        if self.app_icon:
            dialog.iconphoto(False, self.app_icon)
        else:
            # 策略B: 如果主程序用的是 ICO，或者 app_icon 为空，尝试手动加载 ICO
            try:
                ico_path = self.get_resource_path("logo.ico")
                if os.path.exists(ico_path):
                    dialog.iconbitmap(ico_path)
            except Exception:
                pass
        
        # 2. 获取当前主题颜色
        mode = self.theme_mode
        colors = self.colors[mode]
        current_bg = colors["card"]
        current_fg = colors["text"]
        current_input_bg = colors["input_bg"]
        
        dialog.configure(bg=current_bg)
        
        # 3. Windows 标题栏颜色适配 (DWM)
        def force_dark_title_bar(window):
            try:
                import ctypes
                from ctypes import windll, c_int, byref
                if os.name == 'nt':
                    window.update()
                    hwnd = windll.user32.GetParent(window.winfo_id())
                    # 如果没有 transient，GetParent 可能获取不到正确的句柄，尝试直接用 winfo_id
                    if not hwnd: 
                        hwnd = window.winfo_id()
                        # 对于非 transient 窗口，可能需要获取自身的 HWND
                        # 但 Tkinter 的 winfo_id 往往只是 client area
                        # 再次尝试获取当前活动窗口句柄作为兜底
                        # hwnd = windll.user32.GetForegroundWindow() 
                    
                    # 重新获取正确的顶级窗口句柄
                    hwnd = windll.user32.GetParent(window.winfo_id())
                    
                    val = c_int(1)
                    windll.dwmapi.DwmSetWindowAttribute(hwnd, 20, byref(val), 4)
                    windll.dwmapi.DwmSetWindowAttribute(hwnd, 19, byref(val), 4)
            except Exception:
                pass

        if mode == "dark":
            dialog.after(10, lambda: force_dark_title_bar(dialog))

        # 4. 尺寸控制 (主窗口 1/2)
        main_w = self.root.winfo_width()
        main_h = self.root.winfo_height()
        w = max(int(main_w / 2), 500)
        h = max(int(main_h / 2), 300)
        
        root_x = self.root.winfo_rootx()
        root_y = self.root.winfo_rooty()
        x = root_x + (main_w - w) // 2
        y = root_y + (main_h - h) // 2
        
        dialog.geometry(f"{w}x{h}+{x}+{y}")
        dialog.resizable(True, True) 
        
        # 5. 模态控制 (替代 transient)
        dialog.lift()        # 提升到顶层
        dialog.focus_force() # 强制获取焦点
        dialog.grab_set()    # 独占事件（实现模态）
        
        # 6. 布局容器
        main_container = tk.Frame(dialog, bg=current_bg)
        main_container.pack(fill="both", expand=True, padx=int(30*self.scale_factor), pady=int(30*self.scale_factor))
        
        # 标题
        tk.Label(main_container, text="请输入新的API Token:", 
                bg=current_bg, fg=current_fg,
                font=(self.font_family, int(self.title_font_size * 1.2 * self.scale_factor))).pack(anchor="w", pady=(0, 20))
        
        # 输入区域
        entry_frame = tk.Frame(main_container, bg=current_bg)
        entry_frame.pack(fill="x", expand=True, pady=(0, 20))
        
        token_var = tk.StringVar()
        
        entry = tk.Entry(entry_frame, textvariable=token_var, show="•", 
                        font=("Consolas", int(self.base_font_size * 1.2 * self.scale_factor)), 
                        bg=current_input_bg, 
                        fg=current_fg,       
                        insertbackground=current_fg, 
                        relief="solid", bd=1)
        entry.pack(side=tk.LEFT, fill="x", expand=True, padx=(0, 15), ipady=8)
        entry.focus_set()
        
        # 显示/隐藏按钮
        show_var = tk.BooleanVar(value=False)
        def toggle_show():
            if show_var.get():
                entry.config(show="")
                show_btn.config(text="👁 隐藏")
            else:
                entry.config(show="•")
                show_btn.config(text="👁 显示")
        
        show_btn = ModernButton(entry_frame, text="👁 显示", command=toggle_show,
                               variant="secondary", padx=15, pady=8)
        show_btn.pack(side=tk.LEFT)
        
        # 底部按钮区
        btn_frame = tk.Frame(main_container, bg=current_bg)
        btn_frame.pack(side="bottom", fill="x", pady=(20, 0))
        
        def ok():
            new_token = token_var.get().strip()
            if new_token:
                self.token = new_token
                self.token_label.config(text="•" * 20)
                self.log_message("✅ Token 更新成功", "成功")
                self.save_token_to_config(new_token)
                messagebox.showinfo("提示", "Token 已更新")
            dialog.destroy()
        
        ok_btn = ModernButton(btn_frame, text="确定更新", variant="primary",
                             command=ok, width=12, padx=10, pady=5)
        ok_btn.pack(side="right")
        
        cancel_btn = ModernButton(btn_frame, text="取消", variant="secondary",
                                 command=dialog.destroy, width=10, padx=10, pady=5)
        cancel_btn.pack(side="right", padx=(0, 15))
        
        dialog.bind("<Return>", lambda e: ok())
        dialog.bind("<Escape>", lambda e: dialog.destroy())

    def save_token_to_config(self, token):
        config_file = os.path.join(self.get_app_path(), "config.ini")
        config = configparser.ConfigParser()
        
        if os.path.exists(config_file):
            config.read(config_file, encoding='utf-8')
        
        if 'API' not in config:
            config['API'] = {}
        
        config['API']['token'] = token
        
        with open(config_file, 'w', encoding='utf-8') as f:
            config.write(f)
        
        self.log_message(f"🔐 Token 已保存到配置文件: {config_file}", "信息")

    def call_ocr_api_with_retry(self, file_path, retries=3, timeout=(30, 60)): 
        for attempt in range(retries):
            if not self.processing:
                return None
            
            try:
                result = self.call_ocr_api(file_path, timeout)
                
                if not self.processing:
                    return None
                    
                if result is not None:
                    return result
                
                if attempt < retries - 1:
                    wait_time = 2 * (attempt + 1)
                    self.log_message(f"⚠️ 连接不稳，{wait_time}秒后重试 ({attempt+1}/{retries})...", "警告")
                    
                    for _ in range(wait_time * 2): 
                        if not self.processing: return None
                        time.sleep(0.5)
                    
            except Exception as e:
                if attempt < retries - 1:
                    wait_time = 2 * (attempt + 1)
                    self.log_message(f"⏳ 异常重试: {e}", "警告")
                    time.sleep(wait_time)
                else:
                    raise e
        return None
    
    def call_ocr_api(self, file_path, timeout=(60, 300)):
        api_name = self.api_var.get()
        url = self.api_configs[api_name]["url"]
        
        try:
            with open(file_path, "rb") as f:
                data = base64.b64encode(f.read()).decode()
            
            payload = {
                "file": data,
                "fileType": 0 if file_path.lower().endswith('.pdf') else 1
            }
            
            if api_name == "PP-OCRv5":
                 payload["use_doc_preprocessor"] = True
                 payload["use_textline_orientation"] = True
            
            resp = requests.post(
                url,
                json=payload,
                headers={"Authorization": f"token {self.token}", "Connection": "close"},
                timeout=timeout 
            )
            
            if resp.status_code == 200:
                return resp.json()
            elif resp.status_code == 503:
                self.log_message(f"⚠️ 服务暂时不可用 (503)，准备重试...", "警告")
                return None
            
            self.log_message(f"❌ API响应错误 {resp.status_code}: {resp.text[:100]}", "错误")
            return None
            
        except requests.exceptions.Timeout:
            self.log_message(f"⏳ 上传/处理超时 (超过{timeout[1]}秒)", "错误")
            return None
        except requests.exceptions.ConnectionError:
            self.log_message(f"🔌 网络连接断开", "错误")
            return None
        except Exception as e:
            self.log_message(f"❌ API请求异常: {str(e)}", "错误")
            return None
    
    def clean_html_content(self, content):
        content = re.sub(r'<[^>]+>', '', content)
        content = re.sub(r'&nbsp;', ' ', content)
        content = re.sub(r'&lt;', '<', content)
        content = re.sub(r'&gt;', '>', content)
        content = re.sub(r'&amp;', '&', content)
        content = re.sub(r'&quot;', '"', content)
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = re.sub(r'[ \t]+', ' ', content)
        
        return content.strip()
    
    def clear_log(self):
        self.log_text.delete(1.0, tk.END)
        self.log_message("📜 日志已清空", "信息")

if __name__ == "__main__":
    try:
        import ctypes
        myappid = 'mycompany.ocr.desktop.final.v6' 
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    except Exception:
        pass

    root = tk.Tk() 
    root.withdraw() 
    app = OCRDesktopApp(root)
    root.deiconify() 
    root.mainloop()